# Copyright (C) 2015-2020 Lucas Aimaretto / laimaretto@gmail.com
#
# This is taskAutom
#
# taskAutom is free software: you can redistribute it and/or modify
# it under the terms of the 3-clause BSD License.
#
# taskAutom is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY of any kind whatsoever.
#

import paramiko
from paramiko import client
from sshtunnel import SSHTunnelForwarder
from netmiko import ConnectHandler
from scp import SCPClient
import pandas as pd

import docx
from docx.enum.style import WD_STYLE_TYPE 
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
#from traitlets.traitlets import default

import yaml
import sys
import telnetlib
import ftplib
import os
import csv
import time
import threading
from multiprocessing.pool import ThreadPool
from operator import itemgetter
from itertools import groupby
import logging
import importlib
import re
import argparse
from getpass import getpass
import re
import calendar
import random
from socket import timeout

#logging.basicConfig(level=logging.DEBUG,format='[%(levelname)s] (%(threadName)-10s) %(message)s')


# logging.basicConfig(filename='debug.log', level=logging.DEBUG)
# logger = logging.getLogger("netmiko")



# Variables Login
IP_LOCALHOST          	 = "127.0.0.1"

ROUTER_TELNET_PORT       = 23
ROUTER_SSH_PORT          = 22
ROUTER_FTP_PORT          = 21

# --- General Timers
ALU_TIME_LOGIN           = 5
SAM_TIME_LOGIN           = 10
ALU_TIME_DIFF			 = 1
PROMPT_TIMEOUT           = ALU_TIME_LOGIN

# --- General Prompts
ALU_PROMPT_CLOSED         = [b"closed by foreign host"]
ALU_PROMPT_LOGOUT		  = [b"# logout"]
ALU_PROMPT_FTP_LOGOUT     = [b"221 Bye!"]
ALU_PROMPT_LOGIN          = [b"Login:"]
ALU_PROMPT_FTP_LOGIN      = [b"220 FTP server ready"]
ALU_PROMPT_FTP_BIN_MODE   = [b"binary mode"]
ALU_PROMPT_FTP_TXFER      = [b"226 Transfer complete"]
ALU_PROMPT_FTP            = [b"ftp>"]

ALU_PROMPT_PASS           = [b"Password:"]
ALU_PROMPT                = [b"(A:|B:)(.+)(>|#)"]

ALU_TIMOS_LOGIN           = [b"(TiMOS-[A-Z]-\d{1,2}.\d{1,2}.R\d{1,2})"]
ALU_HOSTNAME              = [b"(A:|B:)(.+)(>|#)"]

# --- Extras
CH_CR					  = "\n"
CH_COMA 				  = ","
LOG_GLOBAL                = []

# - Parameters per vendor
DICT_VENDOR = dict(
	nokia_sros=dict(
		START_SCRIPT   = "# SCRIPT_NONO_START", # no \n in the end
		FIRST_LINE     = "\n/environment no more\n",
		LAST_LINE      = "\nexit all\n",
		FIN_SCRIPT     = "# SCRIPT_NONO_FIN", # no \n in the end
		VERSION 	   = "show version", # no \n in the end
		VERSION_REGEX  = "(TiMOS-[A-Z]-\d{1,2}.\d{1,2}.R\d{1,2})",
		HOSTNAME_REGEX = "(A:|B:)(.+)(>|#)",
		SHOW_REGEX     = "(\/show|show)\s.+",
		MAJOR_ERROR_LIST = [b"FAILED:",b"invalid token",b"ERROR:",b"not allowed",b"Error"],
		MINOR_ERROR_LIST = [b"MINOR:"],
	),
)

####

def fncPrintResults(routers, timeTotalStart, dictParam, DIRECTORY_LOG_INFO='', ALU_FILE_OUT_CSV=''):

	separator = "\n------ * ------"

	outTxt    = ""

	outTxt = outTxt + separator + '\n'

	#### GLOBALS

	outTxt = outTxt + "Global Parameters:\n"

	outTxt = outTxt + "  Template File:              " + dictParam['pyFile'] + '\n'
	outTxt = outTxt + "  Template Type:              " + dictParam['pluginType'] + '\n'
	outTxt = outTxt + "  CSV File:                   " + dictParam['csvFile'] + '\n'
	outTxt = outTxt + "  Text File:                  " + "job0_" + dictParam['pyFile'] + ".txt" + '\n'

	if dictParam['genMop'] == 'yes':
		outTxt = outTxt + "  MOP filename                " + "job0_" + dictParam['pyFile'] + ".docx\n"

	outTxt = outTxt + "  Inventory file              " + str(dictParam['inventoryFile']) + "\n"


	outTxt = outTxt + "  Verify Commands:            " + dictParam['cmdVerify'] + '\n'	
	outTxt = outTxt + "  Strict Order:               " + dictParam['strictOrder'] + '\n'

	if dictParam['strictOrder'] == 'yes':
		outTxt = outTxt + "  Halt-on-Error:              " + dictParam['haltOnError'] + '\n'

	if dictParam['logInfo']:
		outTxt = outTxt + "  Folder logInfo:             " + dictParam['logInfo'] + '\n'
	else:
		outTxt = outTxt + "  Folder logInfo:             " + "None" + '\n'

	if len(dictParam['cronTime']):
		outTxt = outTxt + "  CRON Config:                " + str(dictParam['cronTime']) + '\n'
	else:
		outTxt = outTxt + "  CRON Config:                " + "None" + '\n'

	outTxt = outTxt + "  Total Threads:              " + str(dictParam['progNumThreads']) + '\n'

	if dictParam['strictOrder'] == 'no':
		outTxt = outTxt + "  Total Routers:              " + str(len(routers)) + '\n'
	else:
		outTxt = outTxt + "  Total Lines:                " + str(len(routers)) + '\n'

	#### CONNECTION

	outTxt = outTxt + "\nDefault Connection Parameters:\n"

	if dictParam['inventoryFile'] != None:
		outTxt = outTxt + "(Override by inventory file: " + dictParam['inventoryFile'] + ")\n\n"
	
	if dictParam['useSSHTunnel'] == 'yes':
		outTxt = outTxt + "  Use SSH tunnel:             " + str(dictParam['useSSHTunnel']) +" ("+ str(len(dictParam['jumpHosts'])) +")" + '\n'
	else:
		outTxt = outTxt + "  Use SSH tunnel:             " + str(dictParam['useSSHTunnel']) + '\n'
	
	outTxt = outTxt + "  Client Type:                " + str(dictParam['clientType']) + '\n'
	outTxt = outTxt + "  Telnet Timeout:             " + str(dictParam['telnetTimeout']) + "s" + '\n'
	outTxt = outTxt + "  SSH Delay Factor:           " + str(dictParam['delayFactor']) + '\n'
	outTxt = outTxt + "  SSH MaxLoops:               " + str(dictParam['sshMaxLoops']) + '\n'
	outTxt = outTxt + "  Username:                   " + str(dictParam['username']) + '\n'
	outTxt = outTxt + "  Device Type:                " + str(dictParam['deviceType']) + '\n'

	if dictParam['outputJob'] > 0:

		timeTotalEnd 	= time.time()
		timeTotal 		= timeTotalEnd - timeTotalStart		

		outTxt = outTxt + separator + '\n'

		routers = LOG_GLOBAL
		columns=['DateTime','logInfo','Plugin','pluginType','cmdVerify','IP','HostName','User','Reason','id','port','jumpHost','clientType','txLines','rxLines','time','telnetTimeout','delayFactor','sshMaxLoops','servers']
		df = pd.DataFrame(routers,columns=columns)

		outTxt = outTxt + "\nTiming:\n"


		outTxt = outTxt + "  timeMin                     " + fncFormatTime(df['time'].min()) + "s" + '\n'
		outTxt = outTxt + "  timeAvg:                    " + fncFormatTime(df['time'].mean()) + "s" + '\n'
		outTxt = outTxt + "  timeMax:                    " + fncFormatTime(df['time'].max()) + "s" + '\n'
		outTxt = outTxt + "  timeTotal:                  " + fncFormatTime(timeTotal) + "s" + '\n'
		outTxt = outTxt + "  timeTotal/totalRouters:     " + fncFormatTime(timeTotal/len(routers)) + "s" + '\n'

		outTxt = outTxt + separator + '\n'

		df['threads']     = dictParam['progNumThreads']

		df.to_csv(ALU_FILE_OUT_CSV,index=False)

		dfFailed = df[df['Reason'] != 'SendSuccess']

		if dictParam['strictOrder'] == 'no':
			outTxt = outTxt + "\nFailed routers:             " + str(len(dfFailed)) + '\n'
		else:
			outTxt = outTxt + "\nFailed lines:               " + str(len(dfFailed)) + '\n'

		if dictParam['strictOrder'] == 'yes' and dictParam['haltOnError'] == 'yes' and dictParam['aluLogReason'] not in ['SendSucces','ReadTimeout']:
			outTxt = outTxt + "   --> HaltOnError: " + dictParam['aluLogReason'] + ' <--\n'

		if len(dfFailed) > 0:
			outTxt = outTxt + dfFailed.to_string() + '\n'

		outTxt = outTxt + separator
		dfGroup = df.groupby(['Reason']).agg({'Reason':['count'],'time':['min','max']})

		outTxt = outTxt + '\n' + dfGroup.to_string() + '\n'

		with open(DIRECTORY_LOG_INFO + '00_report.txt','w') as f:
			f.write(outTxt)

	outTxt = outTxt + separator + '\n'

	print(outTxt)

def fncFormatTime(timeFloat):

	move = 100

	return str( float(int(timeFloat*move))/move )

def fncPrintConsole(inText, show=1):
	#logging.debug(inText)
	localtime   = time.localtime()
	if show:
		print(str(time.strftime("%H:%M:%S", localtime)) + "| " + inText)

def run_mi_thread(i, CliLine, ip, dictParam):
	"""[summary]

	Args:
		i ([type]): [description]
		CliLine ([type]): [description]
		ip ([type]): [description]
		dictParam ([dict]): [Dictionary with connection parameters]
	"""
	time.sleep(random.random())
	aluLogReason = myConnection(i, CliLine, ip, dictParam).run()

	return aluLogReason

def sort_order(lista):
	"""[List will be ordered and sorted always by the first field which is the system IP of the router]

	Args:
		lista ([list]): [List of IP system]

	Returns:
		[list]: [Ordered List]
	"""

	lista_sorted 	= sorted(lista, key=itemgetter(0))
	lista_grouped 	= groupby(lista_sorted, key=itemgetter(0))
	a = []
	for i,rou in enumerate(lista_grouped):
		a.append(list(rou[1]))
	return a

def verifyCronTime(cronTime):
	"""[We verify cronTime before moving on]

	Args:
		cronTime ([list]): [list of parameters]

	Returns:
		[list]
	"""

	if cronTime in ['',[],None]:
		return []
	elif len(cronTime)!=6:
		print('Wrong cronTime length. Quitting ...')
		quit()
	else:
		cronName   = str(cronTime[0])
		month      = str(cronTime[1])
		weekday    = str(cronTime[2])
		dayOfMonth = int(cronTime[3])
		hour       = int(cronTime[4])
		minute     = int(cronTime[5])

	if cronName[0] in [str(x) for x in range(0,10)]:
		print('Wrong CRON name. First char cannot be a number. Quitting ...')
		quit()		
	elif not re.compile(r'^[0-9A-Za-z]{1,32}$').search(cronName):
		print('Wrong CRON name. Quitting ...')
		quit()
	
	if month not in [calendar.month_name[x].lower() for x in range(1,13)]:
		print('Wrong month name. Quitting ...')
		quit()

	if weekday not in [calendar.day_name[x].lower() for x in range(0,7)]:
		print('Wrong weekDay name. Quitting ...')
		quit()		

	if dayOfMonth not in list(range(1,32)):
		print('Wrong dayOfMonth value. Quitting ...')
		quit()			

	if hour not in list(range(0,24)):
		print('Wrong hour value. Quitting ...')
		quit()

	if minute not in list(range(0,60)):
		print('Wrong minute value. Quitting ...')
		quit()

	return cronTime

def verifyServers(jumpHostsFile):
	"""We verify the SERVERS dictionary before moving on.

	Args:
		SERVERS ([str]): [Name of the file containing servers information in YML format.]

	Returns:
		[dict]: [Dictionary with servers information]
	"""

	try:
		with open(jumpHostsFile,'r') as f:
			servers = yaml.load(f, Loader=yaml.FullLoader)
	except:
		print("Missing " + jumpHostsFile + " file. Quitting..")
		quit()

	fields = ['name','user','password','ip','port']
	for k in servers.keys():
		for f in fields:
			if f in servers[k].keys():
				if not servers[k][f]:
					print('Missing value for field "' + str(f) + '" in server "' + str(k) + '". Quitting...')
					quit()
			else:
				print('Missing field "' + str(f) + '" in server "' + str(k) + '". Quitting...')
				quit()

	# If before checking is ok, we create a new dictionary with correlative keys for those values...
	return servers

def verifyCsv(csvFile):
	"""[Verify CSV file]

	Args:
		csvFile ([str]): [Name of CSV file]

	Returns:
		[list]: [List of Routers]
	"""

	try:
		if csvFile.split(".")[-1] == "csv":
			iFile 		= open(csvFile,"r")
			csvFile 	= csv.reader(iFile, delimiter=",", quotechar="|")
			routers 	= list(csvFile)
			iFile.close()
		else:
			print("Missing CSV file. Verify extension of the file to be '.csv'. Quitting...")
			quit()
	except:
		print("No CSV file found. Quitting ...")
		quit()

	return routers

def verifyPlugin(pyFile):
	"""[Verifies the plugin template]

	Args:
		pyFile ([str]): [Name of config template]

	Returns:
		[module]: [The module]
	"""

	try:
		if pyFile.split(".")[-1] == "py":
			pyFile = pyFile.split(".")[0]
			#exec ("from " + pyFile + " import construir_cliLine")
			mod = importlib.import_module(pyFile)
			print(mod)
		else:
			print("Missing config file. Verify extension of the file to be '.py'. Quitting...")
			quit()
	except Exception as e:
		print(e)
		print("----\nError importing configFile. Quitting ...")
		quit()

	return mod

def verifyConfigFile(config_file):
	""" This function checks the whole text in order to search for ASCII 
	characters (7bit) since 8bit chars won't allow a proper boot process.
	"""

	charset_allowed = [chr(c) for c in range(128)]

	for i,line in enumerate(config_file.split('\n')):
		for character in line:
			if character not in charset_allowed:
				return i+1, line, character

	return -1,-1

def verifyInventory(inventoryFile, jumpHostsFile):

	columns = ['ip','username','password','clientType','useSSHTunnel','telnetTimeout','delayFactor','sshMaxLoops','jumpHost']

	try:
		df = pd.read_csv(inventoryFile)
	except:
		print("Inventory: The file " + inventoryFile + " was not found. Quitting ...")
		quit()

	for col in df.columns:
		if col not in columns:
			print("Inventory: The field '" + col +  "' is not a valid one. Valids are: " + str(columns) + ". Quitting...")
			quit()

	for col in columns:
		if col not in df.columns:
			print("Inventory: The inventory file ("+inventoryFile+") is missing the field '" + col +  "'. Quitting...")
			quit()			

	df2 = df.copy()
	df2 = df2.fillna("")

	for row in df2.itertuples():

		ip   = row.ip
		jh   = row.jumpHost
		ct   = row.clientType
		tun  = row.useSSHTunnel
		to   = row.telnetTimeout
		dfac = row.delayFactor
		sml  = row.sshMaxLoops

		if tun not in ['yes','no','']:
			print("Inventory: The router " + ip + " is not using a valid sshTunnel option. For default, leave empty. Quitting...")
			quit()			

		if tun == 'yes':

			serversList = list(verifyServers(jumpHostsFile).keys()) + ['']

			if jh not in serversList:
				print("Inventory: The router " + ip + " is using sshtunnel and has not a valid jumpHost. If empty, using default. Available: " + str(serversList) + ". Quitting...")
				quit()

		if ct not in ['ssh','tel','']:
			print("Inventory: The router " + ip + " is not using a valid clientType. For default, leave empty. Quitting...")
			quit()

		if to != '':
			try:
				float(to)
			except:
				print("Inventory: The router " + ip + " has not a valid telnetTimeout. For default, leave empty. Quitting...")
				quit()

		if dfac != '':
			try:
				float(dfac)
			except:
				print("Inventory: The router " + ip + " has not a valid delayFactor. For default, leave empty. Quitting...")
				quit()

		if sml != '':
			try:
				float(sml)
			except:
				print("Inventory: The router " + ip + " has not a valid sshMaxLoops. For default, leave empty. Quitting...")
				quit()				


	df3 = df2.set_index('ip').transpose().to_dict()

	return df3

def renderMop(aluCliLineJob0, pyFile, genMop):
	"""[Generates a MOP based on the CSV and plugin information]

	Args:
		aluCliLineJob0 ([file]): [configLines]
		pyFile ([str]):  [The plugin for this MOP]

	Returns:
		None
	"""

	job0docx = "job0_" + pyFile + ".docx"
	job0text = "job0_" + pyFile + ".txt"

	if genMop == 'yes':

		print("\nGenerating MOP: " + job0docx)
		config = aluCliLineJob0.split('\n')
		config = [x for x in config if len(x) > 0]

		myDoc = docx.Document()
		myStyles = myDoc.styles  

		styleConsole = myStyles.add_style('Console', WD_STYLE_TYPE.PARAGRAPH)
		styleConsole.font.name = 'Courier'
		styleConsole.font.size = Pt(9)
		styleConsole.paragraph_format.keep_together = True

		styleConsole.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		#styleConsole.paragraph_format.line_spacing = Pt(10)
		#styleConsole.paragraph_format.line_spacing = .2
		styleConsole.paragraph_format.space_after = Pt(2)

		myDoc.add_heading('MOP for ' + pyFile, 0)

		for i,row in enumerate(config):

			if i == 0:
				myDoc.add_heading('Configuraciones',1)

			if 'Heading_2' in row.split(":")[0]:
				row = ''.join(row.split(":")[1:])
				subtitle = myDoc.add_paragraph(row)
				subtitle.style = myDoc.styles['Heading 2']
				subtitle.paragraph_format.line_spacing = 1.5

			elif 'Heading_3' in row.split(":")[0]:
				row = ''.join(row.split(":")[1:])
				subtitle = myDoc.add_paragraph(row)
				subtitle.style = myDoc.styles['Heading 3']
				subtitle.paragraph_format.line_spacing = 1.5

			else:
				configText = myDoc.add_paragraph(row)
				configText.style = myDoc.styles['Console']

		myDoc.save(job0docx)
		print("MOP done...")

	with open(job0text,'w') as f:
		f.write(aluCliLineJob0)

def renderCliLine(router, dictParam, mod):

	# if strictOrder == yes, the received router vector will be ..
	# [ip1, par1, par2, ... , parN]
	# [ip1, par1, par2, ... , parN]
	# [ip2, par1, par2, ... , parN]
	# [ip2, par1, par2, ... , parN]

	# if strictOrder == no, the received router vector will be ..
	# [[ip1, par1, par2, ... , parN],[ip1, par1, par2, ... , parN]]
	# [[ip2, par1, par2, ... , parN],[ip2, par1, par2, ... , parN]]

	aluCliLine = ""

	if dictParam['outputJob'] == 2:
		mop = None
	else:
		mop = 1

	if dictParam['strictOrder'] == 'no':
		systemIP = router[0][0]
	else:
		systemIP = router[0]

	if dictParam['strictOrder'] == 'no':
		for j,item in enumerate(router):
			aluCliLine = aluCliLine + mod.construir_cliLine(j,item, mop)
	else:
		aluCliLine = mod.construir_cliLine(0,router, mop)

	if aluCliLine[-1] == "\n":
		aluCliLine = aluCliLine[:-1]

	if dictParam['outputJob'] == 2:	

		if len(dictParam['cronTime']) == 0:
			
			pass

		return systemIP, aluCliLine

	elif dictParam['outputJob'] == 0:

		return aluCliLine
###

class myConnection(threading.Thread):
	"""
	[Class for connection Object]
	"""

	def __init__(self, thrdNum, config_line, systemIP, dictParam):

		threading.Thread.__init__(self)
		self.num 			  = thrdNum
		self.datos 			  = DICT_VENDOR[dictParam['deviceType']]['START_SCRIPT'] + DICT_VENDOR[dictParam['deviceType']]['FIRST_LINE'] + config_line + DICT_VENDOR[dictParam['deviceType']]['LAST_LINE'] + DICT_VENDOR[dictParam['deviceType']]['FIN_SCRIPT']
		self.outputJob 	      = dictParam['outputJob']
		self.DIRECTORY_LOGS   = dictParam['DIRECTORY_LOGS']
		self.telnetTimeout    = dictParam['telnetTimeout']
		self.ALU_FILE_OUT_CSV = dictParam['ALU_FILE_OUT_CSV']
		self.delayFactor      = dictParam['delayFactor']
		self.logInfo          = dictParam['logInfo']
		self.LOG_TIME         = dictParam['LOG_TIME']
		self.plugin           = dictParam['pyFile']

		# local generated variables
		self.connInfo = {
			'systemIP':systemIP,
			'useSSHTunnel':dictParam['useSSHTunnel'],
			'clientType':dictParam['clientType'],
			'localPort':-1,
			'remotePort':-1,
			'controlPlaneAccess':-1,
			'aluLogged':-1,
			'username':dictParam['username'],
			'password':dictParam['password'],
			'aluLogReason':"N/A",
			'hostname':"N/A",
			'timos':"N/A",
			'cronTime':dictParam['cronTime'],
			'sshServer':-1,
			'conn2rtr':-1,
			'delayFactor':dictParam['delayFactor'],
			'telnetTimeout':dictParam['telnetTimeout'],
			'telnetWriteTimeout':dictParam['telnetWriteTimeout'],
			'sshMaxLoops':dictParam['sshMaxLoops'],
			'jumpHosts':dictParam['jumpHosts'],
			'inventory':dictParam['inventory'],
			'strictOrder':dictParam['strictOrder'],
			'deviceType':dictParam['deviceType'],
			'pluginType':dictParam['pluginType'],
			'cmdVerify':dictParam['cmdVerify'],
		}

		# Do we you use jumpHosts?
		if self.connInfo['useSSHTunnel'] == 'yes' or dictParam['inventoryFile'] != None:
			self.connInfo['jumpHost'] = [x for i,x in enumerate(self.connInfo['jumpHosts']) if self.num % len(self.connInfo['jumpHosts']) == i][0]
		else:
			self.connInfo['jumpHost'] = -1

		# ### Update per router data with informatino from inventory
		if dictParam['inventoryFile'] != None and self.connInfo['systemIP'] in self.connInfo['inventory'].keys():
			self.tempDict = self.connInfo['inventory'][systemIP]
			for key in self.tempDict.keys():
				if self.tempDict[key] != '':
					self.connInfo[key] = self.tempDict[key]

		# Identify connection ports
		if ":" in self.connInfo['systemIP']:
			self.connInfo['remotePort'] = int( self.connInfo['systemIP'].split(":")[1] )			
			self.connInfo['systemIP']   = self.connInfo['systemIP'].split(":")[0]
		else:
			if self.connInfo['clientType'] == 'tel':
				self.connInfo['remotePort'] = ROUTER_TELNET_PORT

			elif self.connInfo['clientType'] == 'ssh':
				self.connInfo['remotePort'] = ROUTER_SSH_PORT

		# --- Users
		self.ROUTER_USER1    = [self.connInfo['username'],self.connInfo['password']]
		self.ROUTER_USER2    = ["extraUser1","extraPassword1"]
		self.ROUTER_USER3    = ["extraUser2","extraPassword2"]
		self.ROUTER_USER     = [self.ROUTER_USER1]

		self.tDiff	    = 0
		self.strConn    = "Con-" + str(self.num) + "| "
		self.outRx 	    = ''
		self.fRx        = ''
		self.runStatus  = 1
		self.useCron    = len(self.connInfo['cronTime'])
		
	def run(self):

		# We update the connection info dictionary, after we've set up the connection towards the router...
		self.connInfo.update(self.fncConnectToRouter(self.connInfo))

		if self.connInfo['conn2rtr'] != -1 and self.connInfo['aluLogged'] == 1:
			
			self.connInfo['timos']      = self.fncAuxGetVal(self.connInfo, 'timos')
			self.connInfo['hostname']   = self.fncAuxGetVal(self.connInfo, 'hostname')
			self.connInfo['timosMajor'] = int(self.connInfo['timos'].split("-")[2].split(".")[0])
			
			if self.outputJob == 2:

				fncPrintConsole(self.strConn + "#### Running routine for " + self.connInfo['systemIP'] +  " ...")

				self.f = self.logFileCreation(self.connInfo['hostname'], self.DIRECTORY_LOGS, self.datos, self.strConn)

				self.fRx		 = self.f[0]
				self.fullPathCmd = self.f[1]
				self.fCmd        = self.f[2]

				if self.useCron > 0:

					self.s = self.fncUploadFile(self.strConn, self.fullPathCmd, self.fCmd, self.connInfo)

					self.sftpStatus   = self.s[0]
					self.aluLogReason = self.s[1]

					if self.sftpStatus == 1:

						self.datos = self.runCron(self.fCmd, self.connInfo)
						self.b     = self.routerRunRoutine(self.datos, self.connInfo)

						#fncPrintConsole(self.strConn + "Run: " + str(self.b[0]))

						self.connInfo['aluLogReason'] = self.b[0]
						self.tDiff 					  = self.b[1]
						self.runStatus      		  = self.b[2]
						self.outRx          		  = self.b[3]

				else:
					
					self.b = self.routerRunRoutine(self.datos, self.connInfo)
	
					self.connInfo['aluLogReason'] = self.b[0]
					self.tDiff 					  = self.b[1]
					self.runStatus      		  = self.b[2]
					self.outRx          		  = self.b[3]

				if self.runStatus == 1:

					self.connInfo.update(self.routerLogout(self.connInfo))
					fncPrintConsole(self.strConn + "Logout: " + str(self.connInfo['aluLogReason']))

				else:

					fncPrintConsole(self.strConn + str(self.connInfo['aluLogReason']))

		self.logData(self.connInfo, self.num, self.tDiff, self.outRx, self.fRx, self.strConn, self.datos, self.logInfo, self.LOG_TIME, self.plugin)

		#######################
		# closing connections #

		#print(self.connInfo['conn2rtr'], self.connInfo['aluLogged'], self.connInfo['useSSHTunnel'], self.connInfo['sshServer'].tunnel_is_up, self.connInfo['clientType'])
		if self.connInfo['conn2rtr'] != -1 or self.connInfo['aluLogged'] == 1:

			if self.connInfo['clientType'] == 'tel':
				self.connInfo['conn2rtr'].close()

			elif self.connInfo['clientType'] == 'ssh':
				self.connInfo['conn2rtr'].disconnect()

		if self.connInfo['useSSHTunnel'] == 'yes' and self.connInfo['sshServer']:
			self.connInfo['sshServer'].stop()

		#                     #
		#######################

		return self.connInfo['aluLogReason']

	def fncWriteToConnection(self, inText, connInfo):

		clientType         = connInfo['clientType']
		conn2rtr           = connInfo['conn2rtr']
		telnetWriteTimeout = connInfo['telnetWriteTimeout']
		pluginType         = connInfo['pluginType']

		maxLoops           = connInfo['sshMaxLoops']
		delayFactor        = connInfo['delayFactor']

		if connInfo['cmdVerify'] == 'yes':
			cmdVerify = True
		else:
			cmdVerify = False

		### Writes to a connection. 
		# For telnet connections, stream needs to be encoded before doing it...
		if clientType == 'tel':

			inText = inText + '\n'
			output = conn2rtr.write(inText.encode())
			time.sleep(telnetWriteTimeout)
			cmdType = "genericTelnet"

		# SSH Connections
		elif clientType == 'ssh':

			if type(inText) == type([]):

				if pluginType == 'config':
					output = conn2rtr.send_config_set(config_commands=inText, enter_config_mode=False, cmd_verify=cmdVerify, delay_factor=delayFactor, max_loops=maxLoops)
				elif pluginType == 'show':
					output = ''
					for cmd in inText:
						output = output + '\n' + cmd + '\n' + conn2rtr.send_command(cmd, delay_factor=delayFactor, max_loops=maxLoops)
			
			elif type(inText) == type(''):

				output = conn2rtr.send_command(inText, delay_factor=delayFactor, max_loops=maxLoops)

			return output

	def fncAuxGetVal(self, connInfo, what):

		if connInfo['clientType'] == 'tel':

			if what == "timos":

				inText = "show version\n"	
				self.fncWriteToConnection(inText, connInfo)
				rx     = connInfo['conn2rtr'].expect(ALU_TIMOS_LOGIN)
				timos  = rx[1].groups()[0].decode()

				return timos

			elif what == "hostname":

				inText = "\n"
				self.fncWriteToConnection(inText, connInfo)
				rx     = connInfo['conn2rtr'].expect(ALU_HOSTNAME)
				hostname = rx[1].groups()[1].decode()

				return hostname	

		elif connInfo['clientType'] == 'ssh':

			if what == "timos":

				inText  = DICT_VENDOR[connInfo['deviceType']]['VERSION']
				rx      = self.fncWriteToConnection(inText, connInfo)
				inRegex = DICT_VENDOR[connInfo['deviceType']]['VERSION_REGEX']
				match   = re.compile(inRegex).search(rx)
				try:
					timos   = match.groups()[0]
				except:
					timos   = "not-matched"

				return timos

			elif what == 'hostname':

				inRegex  = DICT_VENDOR[connInfo['deviceType']]['HOSTNAME_REGEX']
				newHn    = connInfo['conn2rtr'].find_prompt()
				match    = re.compile(inRegex).search(newHn)

				try:
					hostname = match.groups()[1]
				except:
					hostname = "not-matched"

				return hostname

	def fncConnectToRouter(self, connInfo):
		"""[We update the connection info dictionary, after we've set up the connection towards the router]

		Args:
			connInfo ([dict]): [Contains all conection related relevant information ]

		Returns:
			[dict]: [Updated connInfo dictionary]
		"""

		### SSH Tunnel

		if connInfo['useSSHTunnel'] == 'yes':

			tunnel = self.fncSshServer(self.strConn, connInfo)

			connInfo['controlPlaneAccess'] 	= tunnel[0]
			connInfo['localPort'] 		   	= tunnel[1]
			connInfo['sshServer']    		= tunnel[2]

			fncPrintConsole(self.strConn + "Trying router " + IP_LOCALHOST + ":" + str(connInfo['localPort']) + " -> " + connInfo['systemIP'] + ":" + str(connInfo['remotePort']))

		else:

			fncPrintConsole(self.strConn + "Using direct " + connInfo['clientType'] + " access: ")
			fncPrintConsole(self.strConn + "Trying router " + connInfo['systemIP'] + ":" + str(connInfo['remotePort']) )

			connInfo['controlPlaneAccess'] 	= 1	
			connInfo['localPort'] 			= connInfo['remotePort']
			connInfo['sshServer']    		= -1

		### Connect to router

		if connInfo['controlPlaneAccess'] == 1:

			if connInfo['clientType'] == 'tel':

				try:
					if connInfo['useSSHTunnel'] == 'yes':
						connInfo['conn2rtr'] = telnetlib.Telnet(IP_LOCALHOST, connInfo['localPort'])
					else:
						connInfo['conn2rtr'] = telnetlib.Telnet(connInfo['systemIP'], connInfo['remotePort'])

					connInfo['conn2rtr'].timeout = ALU_TIME_LOGIN
					a = self.routerLoginTelnet(connInfo)

					connInfo['aluLogged']    = a[0]
					connInfo['username']     = a[1]
					connInfo['aluLogReason'] = a[2]
					connInfo['password']     = a[3]

				except:

					connInfo['conn2rtr'] = -1

			elif connInfo['clientType'] == 'ssh':

				try:
					a = self.routerLoginSsh(connInfo)

					connInfo['conn2rtr']     = a[0]
					connInfo['aluLogged']    = a[1]
					connInfo['username']     = a[2]
					connInfo['aluLogReason'] = a[3]
					connInfo['password']     = a[4]

				except:
					connInfo['conn2rtr'] = -1

		else:

			connInfo['conn2rtr']     = -1
			connInfo['aluLogged'] 	 = -1
			connInfo['username']     = "N/A"
			connInfo['aluLogReason'] = "noControlPlaneAccess"
			connInfo['password']     = "N/A"
			connInfo['sshServer']    = -1

		return connInfo

	def fncUploadFile(self, strConn, fileLocal, fileRemote, connInfo):
		### upload configFile via SFTP

		out = [-1,'sftpError']

		if connInfo['useSSHTunnel'] == 'yes':

			# We need to rewrite the remotePort because the clientType could be telnet.
			# There is no problem because the connection to the CLI has already been 
			# established and is located in connInfo['conn2rtr'].
			# This is only for the purpose of uploading a file vía SFTP.

			connInfo['remotePort'] = ROUTER_SSH_PORT

			sshSftp = self.fncSshServer(strConn, connInfo)
			sftpAccess    = sshSftp[0]
			sftpPort      = sshSftp[1]
			sshServerSftp = sshSftp[2]

			transport = paramiko.Transport((IP_LOCALHOST,sftpPort))
			transport.connect(None,connInfo['username'],connInfo['password'])

			# The routers with timos above 6.X do support SFTP.
			# Otherwise we need to use SCP.

			if connInfo['timosMajor'] > 6:
				fncPrintConsole(strConn + "uploading file: SFTP: " + str(sftpPort))
				sftp = paramiko.SFTPClient.from_transport(transport)
			else:
				fncPrintConsole(strConn + "uploading file: SCP: " + str(sftpPort))
				sftp = SCPClient(transport)

			try:
				sftp.put(fileLocal,'cf3:/' + fileRemote)
				out = [1,'sftpOk']
			except:
				out = [-1,'sftpError']

			sftp.close()
			transport.close()
			sshServerSftp.stop()

		else:

			transport = paramiko.Transport((connInfo['systemIP'],ROUTER_SSH_PORT))
			transport.connect(None,connInfo['username'],connInfo['password'])

			if connInfo['timosMajor'] > 6:
				fncPrintConsole(strConn + "uploading file: SFTP: " + str(sftpPort))
				sftp = paramiko.SFTPClient.from_transport(transport)
			else:
				fncPrintConsole(strConn + "uploading file: SCP: " + str(sftpPort))
				sftp = SCPClient(transport)

			try:
				sftp.put(fileLocal,'cf3:/' + fileRemote)
				out = [1,'sftpOk']
			except:
				out = [-1,'sftpError']

			sftp.close()
			transport.close()

		return out

	def fncSshServer(self, strConn, connInfo):

		#jumpHost = random.choice(list(SERVERS.keys())) 
		jumpHost = connInfo['jumpHost']
		servers  = connInfo['jumpHosts']

		tempIp   = servers[jumpHost]['ip']
		tempPort = servers[jumpHost]['port']
		tempUser = servers[jumpHost]['user']
		tempPass = servers[jumpHost]['password']

		try:
			server = SSHTunnelForwarder( 	(tempIp, tempPort), 
												ssh_username = tempUser, 
												ssh_password = tempPass, 
												remote_bind_address = (connInfo['systemIP'], connInfo['remotePort']),
												allow_agent = False,
											)
			server.start()
			localPort = server.local_bind_port
			controlPlaneAccess = 1
			fncPrintConsole(self.strConn + "sshServerTunnel on port: " + str(localPort))

		except Exception as e:

			fncPrintConsole(strConn, e)
			fncPrintConsole(strConn + "Error SSH Tunnel")
			controlPlaneAccess = -1
			localPort 		   = -1

		return controlPlaneAccess, localPort, server

	def routerLoginTelnet(self, connInfo):

		#	i[0]	meaning
		#	-1		Timeout
		#	0		Login:
		#	1		Connection closed by foreign host

		# NOTA: tener presente que 'i' debe esperarse al primer intento.
		# Si volvemos, es porque el expect ya nos dio login o timeout

		aluLogged    = -1
		aluLogUser   = "N/A"
		aluPass      = "PassN/A"
		aluLogReason = "N/A"
		index        = 0

		conn2rtr     = connInfo['conn2rtr']
		clientType   = connInfo['clientType']
		ALU_TELNET_WRITE_TIMEOUT = connInfo['telnetWriteTimeout']
		systemIP     = connInfo['systemIP']


		while aluLogged == -1:

			try:
				i = conn2rtr.expect(ALU_PROMPT_LOGIN + ALU_PROMPT_CLOSED + ALU_PROMPT , ALU_TIME_LOGIN)
			except:
				aluLogUser 			= "UserN/A"
				aluLogReason 		= "TelnetError"
				aluLogged 			= -1
				aluPass             = "PassN/A"
				fncPrintConsole(self.strConn + aluLogReason)
				return (aluLogged,aluLogUser,aluLogReason,aluPass)

			# expected: (0, <_sre.SRE_Match object at 0x7f0887a37d98>, 'login:')
			#fncPrintConsole("i: " + str(i))

			if i[0] == -1:
				# timeout
				self.fncWriteToConnection("\003", connInfo)
				aluLogUser 			= "UserN/A"
				aluLogReason 		= "TelnetTimeout"
				aluLogged 			= -1
				fncPrintConsole(self.strConn + aluLogReason)
				break

			elif i[0] == 0:
		
				if index < len(self.ROUTER_USER):

					tempUser = self.ROUTER_USER[index][0]
					tempPass = self.ROUTER_USER[index][1]

					self.fncWriteToConnection(tempUser, connInfo)

					j = conn2rtr.expect(ALU_PROMPT_PASS + ALU_PROMPT_CLOSED)
					# expected: (0, <_sre.SRE_Match object at 0x7f0887a37e00>, ' Password:')
					#fncPrintConsole("j: " + str(j))

					if j[0] == 0:
						self.fncWriteToConnection(tempPass, connInfo)
						#fncPrintConsole(self.strConn + "User: " + tempUser + ", Pass: " + tempPass + ", index: " + str(index))
						aluLogUser 		= tempUser
						aluLogged		= -1
						index			= index + 1


					elif i[0] == 1:
						# ALU_PROMPT_CLOSED
						# Sometimes loggin into a router is not possible
						# because many users are already logged in into it.
						self.fncWriteToConnection("\003", connInfo)
						aluLogUser 			= "UserN/A"
						aluLogReason	 	= "TelnetFailedConnection"
						aluLogged 			= -1
						aluPass             = "PassN/A"
						fncPrintConsole(self.strConn + aluLogReason + ": " + systemIP)
						break

				else:
					# We've tryed all the user/pass. Quitting.
					self.fncWriteToConnection("\003", connInfo)
					aluLogUser 			= tempUser
					aluLogReason	 	= "MaxLoginReached"
					aluLogged 			= -1
					aluPass             = "PassN/A"
					fncPrintConsole(self.strConn + aluLogReason + ": " + systemIP)
					break

			elif i[0] == 1:
				# ALU_PROMPT_CLOSED
				# Sometimes loggin into a router is not possible
				# because many users are already logged in into it.
				self.fncWriteToConnection("\003", connInfo)
				aluLogUser 			= "UserN/A"
				aluLogReason 		= "TelnetFailedConnection"
				aluLogged 			= -1
				aluPass             = "PassN/A"
				fncPrintConsole(self.strConn + aluLogReason + ": " + systemIP)
				break

			elif i[0] > 1:
				aluLogReason 	= "LoggedOk"
				aluLogged 		= 1
				
		return (aluLogged,aluLogUser,aluLogReason,tempPass)

	def routerLoginSsh(self, connInfo):

		conn2rtr     = -1
		aluLogged    = -1
		aluLogUser   = "N/A"
		aluPass      = "PassN/A"
		aluLogReason = "N/A"
		index        = 0

		systemIP     = connInfo['systemIP']
		delayFactor  = connInfo['delayFactor']
		deviceType   = connInfo['deviceType']

		if connInfo['useSSHTunnel'] == 'yes':
			ip   = IP_LOCALHOST
			port = connInfo['localPort']
		else:
			ip   = connInfo['systemIP']
			port = connInfo['remotePort']

		while aluLogged == -1 or index < len(self.ROUTER_USER):

			#if index < len(self.ROUTER_USER):

			tempUser = self.ROUTER_USER[index][0]
			tempPass = self.ROUTER_USER[index][1]
			index 	 = index + 1

			try:
				conn2rtr = ConnectHandler(device_type=deviceType, host=ip, port=port, username=tempUser, password=tempPass, fast_cli = False)
				aluLogged    = 1
				aluLogReason = "LoggedOk"
				aluLogUser   = tempUser
				aluPass      = tempPass
			except Exception as e:
				aluLogUser   = tempUser
				aluLogReason = "SSHFailedConnection"
				aluLogReason = str(e)
				aluLogged 	 = -1
				aluPass      = "PassN/A"
				fncPrintConsole(self.strConn + aluLogReason + ": " + systemIP)

			# else:
			# 	# We've tryed all the user/pass. Quitting.
			# 	aluLogUser 	 = tempUser
			# 	aluLogReason = "MaxLoginReached"
			# 	aluLogged 	 = -1
			# 	aluPass      = "PassN/A"
			# 	fncPrintConsole(self.strConn + aluLogReason + ": " + systemIP)
			# 	break
				
		return (conn2rtr,aluLogged,aluLogUser,aluLogReason,tempPass)

	def logFileCreation(self, hostname, DIRECTORY_LOGS, datos, strConn):

		fncPrintConsole(strConn + "Creating files locally for " + hostname + "...")

		# Verify for logs directory
		if not os.path.exists(DIRECTORY_LOGS):
			os.makedirs(DIRECTORY_LOGS)

		# Filenames
		aluFileCommands = hostname + "_commands.cfg"
		aluFileOutRx	= hostname + "_rx.txt"

		# Complete = Directories + Filenames
		aluCompleteCmd 	= DIRECTORY_LOGS + aluFileCommands
		aluCompleteRx	= DIRECTORY_LOGS + aluFileOutRx

		# Create files
		fCmd = open(aluCompleteCmd, "a")
		fCmd.write(datos)
		fCmd.close()

		fRx	= open(aluCompleteRx, "a")

		return(fRx, aluCompleteCmd, aluFileCommands)

	def routerRunRoutine(self, datos, connInfo):

		# Sending script to ALU
		runStatus    = 1
		tStart 		 = time.time()
		outRx  		 = ""
		aluLogReason = ""
		fin_script   = DICT_VENDOR[connInfo['deviceType']]['FIN_SCRIPT']
		major_error_list = DICT_VENDOR[connInfo['deviceType']]['MAJOR_ERROR_LIST']
		minor_error_list = DICT_VENDOR[connInfo['deviceType']]['MINOR_ERROR_LIST']

		if connInfo['cronTime']:
			fncPrintConsole(self.strConn + "Establishing script with CRON...", show=1)
		else:
			# Splitting self.datos into individual lines
			fncPrintConsole(self.strConn + "Running script per line...", show=1)

		try:

			if connInfo['clientType'] == 'tel':		
				self.fncWriteToConnection(datos, connInfo)
				outRx = connInfo['conn2rtr'].read_until(fin_script.encode(), connInfo['telnetTimeout'])
				outRx = outRx.decode()
			elif connInfo['clientType'] == 'ssh':
				datos = datos.split('\n')
				outRx = self.fncWriteToConnection(datos, connInfo)

		except ConnectionResetError:
			aluLogReason = "ConnectionResetError"
			runStatus = -1
		except EOFError as e:
			aluLogReason = "EOFError"
			runStatus = -1
		except Exception as e:
			#aluLogReason = "ConnectionGeneralError"
			aluLogReason = str(e)
			runStatus = -1

		tEnd  = time.time()
		tDiff = tEnd - tStart

		## Analizing output only if writing to connection was successfull
		if aluLogReason == "":

			str_major_error_list = [x.decode() for x in major_error_list]
			str_minor_error_list = [x.decode() for x in minor_error_list]
			
			if fin_script not in outRx:
				aluLogReason = "ReadTimeout"	
				runStatus    = -1
			elif any(word in outRx for word in str_major_error_list):
				aluLogReason = "MajorFailed"
			elif any(word in outRx for word in str_minor_error_list):
				aluLogReason = "MinorFailed"
			else:
				aluLogReason = "SendSuccess"

		fncPrintConsole(self.strConn + "Time: " + fncFormatTime(tDiff) + ". Result: " + aluLogReason, show=1)

		return(aluLogReason, tDiff, runStatus, outRx)

	def logData(self, connInfo, connId, tDiff, outRx, fRx, strConn, datos, logInfo, LOG_TIME, plugin):

		if connInfo['aluLogged'] == 1:
			fRx.write(outRx)
			fRx.close()

		if connInfo['useSSHTunnel'] == 'yes':
			serverName = connInfo['jumpHost']
			lenServers = len(connInfo['jumpHosts'])
		else:
			serverName = '-1'
			lenServers = '-1'

		aluCsvLine = [
			LOG_TIME,
			logInfo,
			plugin,
			connInfo['pluginType'],
			connInfo['cmdVerify'],
			connInfo['systemIP'],
			connInfo['hostname'],
			connInfo['username'],
			connInfo['aluLogReason'],
			str(connId),
			str(connInfo['localPort']),
			serverName,
			connInfo['clientType'],
			str(len(datos.split('\n'))),
			str(len(outRx.split('\n'))),
			float(fncFormatTime(tDiff)),
			str(connInfo['telnetTimeout']),
			str(connInfo['delayFactor']),
			str(connInfo['sshMaxLoops']),
			str(lenServers),
			]

		fncPrintConsole(strConn + "logData: " + str(aluCsvLine))

		LOG_GLOBAL.append(aluCsvLine)

	def routerLogout(self, connInfo):

		if connInfo['aluLogged'] == 1:

			if connInfo['clientType'] == 'tel':

				#
				# aluPrompt = [ "A:.*#" , "A:.*>.*#" , "B:.*#", "B:.*>.*#" ]
				#                 0           1            2         3

				# loggin correct, proceed with logout
				self.fncWriteToConnection(CH_CR, connInfo)
				i = connInfo['conn2rtr'].expect(ALU_PROMPT, PROMPT_TIMEOUT)
				#fncPrintConsole("i: " + str(i))

				if i[0] in [0,1,2,3]:
					# Logging out
					self.fncWriteToConnection("logout", connInfo)
					
					try:
						j = connInfo['conn2rtr'].expect(ALU_PROMPT_LOGOUT, PROMPT_TIMEOUT)
						#fncPrintConsole("j: " + str(j))
					
						if j[0] in [0,1,2,3]:
							#connInfo['aluLogged'] = -1
							fncPrintConsole(self.strConn + "Logged out OK from " + connInfo['systemIP'])

					except:
						#connInfo['aluLogged'] = -1
						fncPrintConsole(self.strConn + "Something happended. Not properly logged out from " + connInfo['systemIP'])

			elif connInfo['clientType'] == 'ssh':

				pass

		return connInfo

	def sshStop(self):
		self.sshServer.stop()
		fncPrintConsole(self.strConn + "SSH" + str(self.num) + " stopped ...")

	def runCron(self, script, connInfo):

		def setScript(cronName, script):

			cfg = ""
			cfg = cfg + "script " + cronName + " owner taskAutom\nshutdown\n"
			cfg = cfg + "location cf3:\\" + script + "\n"
			cfg = cfg + "no shutdown\n"
			cfg = cfg + "exit\n"
			return cfg

		def action(cronName):

			cfg = ""
			cfg = cfg + "action " + cronName + " owner taskAutom\nshutdown\n"
			cfg = cfg + "results cf3:\\resultTestCron.txt\n"
			cfg = cfg + "script " + cronName + " owner taskAutom\n"
			cfg = cfg + "no shutdown\n"
			cfg = cfg + "exit\n"
			return cfg

		def policy(cronName):

			cfg = ""
			cfg = cfg + "script-policy " + cronName + " owner taskAutom\nshutdown\n"
			cfg = cfg + "results cf3:\\resultTestCron.txt\n"
			cfg = cfg + "script " + cronName + " owner taskAutom\n"
			cfg = cfg + "no shutdown\n"
			cfg = cfg + "exit\n"			
			return cfg

		def schedule(timos, cronName, month, weekday, dayOfMonth, hour, minute):

			cfg = ""
			cfg = cfg + "schedule " + cronName + " owner taskAutom\nshutdown\n"

			if timos > 7:
				cfg = cfg + "script-policy " + cronName + " owner taskAutom\n"
			else:
				cfg = cfg + "action " + cronName + " owner taskAutom\n"
			
			cfg = cfg + "type oneshot\n"
			cfg = cfg + "day-of-month " + dayOfMonth + "\n"
			cfg = cfg + "hour " + hour + "\n"
			cfg = cfg + "minute " + minute + "\n"
			cfg = cfg + "month " + month + "\n"
			cfg = cfg + "weekday " + weekday + "\n"
			cfg = cfg + "no shutdown \n"
			cfg = cfg + "exit\n"
			cfg = cfg + "exit all\n"
			cfg = cfg + "admin save\n"
			cfg = cfg + "echo " + ALU_FIN_SCRIPT + "\n"
			return cfg

		cronName   = str(connInfo['cronTime'][0])
		month      = str(connInfo['cronTime'][1])
		weekday    = str(connInfo['cronTime'][2])
		dayOfMonth = str(connInfo['cronTime'][3])
		hour       = str(connInfo['cronTime'][4])
		minute     = str(connInfo['cronTime'][5])

		cfg = ""

		if connInfo['timosMajor'] > 7:

			cfg = cfg + "/configure system script-control\n"
			cfg = cfg + setScript(cronName, script)
			cfg = cfg + policy(cronName)
			cfg = cfg + "/configure system cron\n"
			cfg = cfg + schedule(connInfo['timosMajor'], cronName, month, weekday, dayOfMonth, hour, minute)

		else:

			cfg = cfg + "/configure cron\n"
			cfg = cfg + setScript(cronName, script)
			cfg = cfg + action(cronName)
			cfg = cfg + schedule(connInfo['timosMajor'], cronName, month, weekday, dayOfMonth, hour, minute)

		cfg = "/environment no more\necho " + ALU_START_SCRIPT + "\n" + cfg

		return cfg

####################################
# Main Function                    #
####################################

def fncRun(dictParam):
	"""[summary]

	Args:
		dictParam ([dict]): [Dictionary with parameters for the connections]
	Returns:
		[int]: 0
	"""
	################
	# Checking...

	# CronTime
	dictParam['cronTime'] = verifyCronTime(dictParam['cronTime'])

	# Servers
	dictParam['jumpHosts'] = {}
	if dictParam['useSSHTunnel'] == 'yes' or dictParam['inventoryFile'] != None:
		dictParam['jumpHosts'] = verifyServers(dictParam['jumpHostsFile'])

	# CSV File
	routers = verifyCsv(dictParam['csvFile'])

	# Config File
	mod = verifyPlugin(dictParam['pyFile'])

	# Inventory
	dictParam['inventory'] = {}
	if dictParam['inventoryFile'] != None:
		dictParam['inventory'] = verifyInventory(dictParam['inventoryFile'], dictParam['jumpHostsFile'])

	# Strict Order
	if dictParam['strictOrder'] == 'yes':
		dictParam['progNumThreads'] = 1

	# Parsing Data
	if dictParam['strictOrder'] == 'no':
		routers = sort_order(routers)

	timeTotalStart 	= time.time()

	# Generar threads
	threads_list 	= ThreadPool(dictParam['progNumThreads'])
	global global_lock
	global_lock     = threading.Lock()

	################
	# Running...
	if dictParam['outputJob'] == 2:

		# logInfo
		dictParam['LOG_TIME']         = time.strftime('%Y-%m-%d_%H-%M-%S', time.localtime())
		dictParam['DIRECTORY_LOGS']   = os.getcwd() + "/logs_" + dictParam['LOG_TIME'] + "_" + dictParam['logInfo'] + "_" + dictParam['pyFile'] + "/"
		dictParam['ALU_FILE_OUT_CSV'] = dictParam['DIRECTORY_LOGS'] + "00_log.csv"

		# Verify if DIRECTORY_LOGS exists. If so, ask for different name ...
		if os.path.exists(dictParam['DIRECTORY_LOGS']):
			print("Folder " + dictParam['DIRECTORY_LOGS'] + " already exists.\nUse a different folder name.\nQuitting ...")
			quit()
		else:
			os.makedirs(dictParam['DIRECTORY_LOGS'])
			open(dictParam['ALU_FILE_OUT_CSV'],'w').close()
			#os.mknod(ALU_FILE_OUT_CSV)

		###############
		# Let's run ....
		for i, router in enumerate(routers):

			systemIP, aluCliLine = renderCliLine(router, dictParam, mod)

			# running routine
			if dictParam['strictOrder'] == 'no':
				threads_list.apply_async(run_mi_thread, args=(i, aluCliLine, systemIP, dictParam))
			else:
				aluLogReason = run_mi_thread(i, aluCliLine, systemIP, dictParam)

				if dictParam['haltOnError'] == 'yes' and aluLogReason not in ['SendSuccess','ReadTimeout']:
					dictParam['aluLogReason'] = aluLogReason
					break

		if dictParam['strictOrder'] == 'no':
			threads_list.close()
			### The .join() implies that processes/threads need to finish themselves before moving on.
			threads_list.join()

		print("all done")
		fncPrintResults(routers, timeTotalStart, dictParam, dictParam['DIRECTORY_LOGS'], dictParam['ALU_FILE_OUT_CSV'])

	elif dictParam['outputJob'] == 0:

		aluCliLineJob0  = ""

		for i, router in enumerate(routers):

			aluCliLineJob0 = aluCliLineJob0 + renderCliLine(router, dictParam, mod)

		verif = verifyConfigFile(aluCliLineJob0)

		if verif != (-1,-1):
			print("\nWrong config file for router " + str(router) + "\nCheck (n,line,char): " + str(verif) + "\nQuitting...")
			quit()			

		renderMop(aluCliLineJob0, dictParam['pyFile'], dictParam['genMop'])
		fncPrintResults(routers, timeTotalStart, dictParam)

	return 0

if __name__ == '__main__':

	parser1 = argparse.ArgumentParser(description='Task Automation Parameters.', prog='PROG', usage='%(prog)s [options]')
	parser1.add_argument('-v'  ,'--version',     help='Version', action='version', version='Lucas Aimaretto - (c)2021 - laimaretto@gmail.com - Version: 7.10.1' )

	parser1.add_argument('-j'  ,'--jobType',       type=int, required=True, choices=[0,2], default=0, help='Type of job')
	parser1.add_argument('-csv','--csvFile',       type=str, required=True, help='CSV File with parameters',)
	parser1.add_argument('-py' ,'--pyFile' ,       type=str, required=True, help='PY Template File',)

	parser1.add_argument('-u'  ,'--username',      type=str, help='Username', )
	parser1.add_argument('-th' ,'--threads' ,      type=int, help='Number of threads. Default=1', default=1,)
	parser1.add_argument('-log','--logInfo' ,      type=str, help='Description for log folder', )
	parser1.add_argument('-jh' ,'--jumpHostsFile', type=str, help='jumpHosts file. Default=servers.yml', default='servers.yml')
	parser1.add_argument('-inv','--inventoryFile', type=str, help='inventory.csv file with per router connection parameters. Default=None', default=None)
	parser1.add_argument('-pt' ,'--pluginType',    type=str, help='Type of plugin. Default=config', default='config', choices=['show','config'])
	parser1.add_argument('-gm', '--genMop',        type=str, help='Generate MOP. Default=no', default='no', choices=['no','yes'])	
	parser1.add_argument('-crt','--cronTime',      type=str, nargs='+' , help='Data for CRON: name(ie: test), month(ie april), weekday(ie monday), day-of-month(ie 28), hour(ie 17), minute(ie 45).', default=[])
	parser1.add_argument('-to' ,'--telnetTimeout', type=int, help='Telnet read Timeout [sec]. Default=90', default=90,)
	parser1.add_argument('-tw' ,'--telnetWriteTimeout', type=int, help='Telnet write Timeout [sec]. DO NOT MODIFY. Default=0', default=0,)
	parser1.add_argument('-df' ,'--delayFactor',   type=float, help='SSH delay factor. Increase if the network is lossy and/on noissy. Improves interaction with the network. Default=1', default=1,)
	parser1.add_argument('-sml' ,'--sshMaxLoops',    type=float, help='SSH MaxLoops. Increase if long outputs are to be expected per each command (mainly for show commands). Default=5000', default=5000)
	parser1.add_argument('-tun','--sshTunnel',     type=str, help='Use SSH Tunnel to routers. Default=yes', default='yes', choices=['no','yes'])
	parser1.add_argument('-ct', '--clientType',    type=str, help='Connection type. Default=ssh', default='ssh', choices=['tel','ssh'])
	parser1.add_argument('-dt', '--deviceType',    type=str, help='Device Type. Default=nokia_sros', default='nokia_sros', choices=['nokia_sros'])
	parser1.add_argument('-so', '--strictOrder',   type=str, help='Follow strict order of routers inside the csvFile. If enabled, threads = 1. Default=no', default='no', choices=['no','yes'])
	parser1.add_argument('-hoe','--haltOnError',   type=str, help='If using --strictOrder, halts if error found on execution. Default=no', default='no', choices=['no','yes'])
	parser1.add_argument('-cv', '--cmdVerify',     type=str, help='Enable cmdVerify when interacting with router. Disable only if connection problems. Default=yes', default='yes', choices=['no','yes'])

	args = parser1.parse_args()

	### reading parameters

	dictParam = dict(
		outputJob 			= args.jobType,
		csvFile 			= args.csvFile,
		pyFile              = args.pyFile,
		username 			= args.username,
		password 			= None,
		progNumThreads		= args.threads,
		logInfo 			= args.logInfo,
		useSSHTunnel 		= args.sshTunnel,
		telnetTimeout 		= args.telnetTimeout,
		telnetWriteTimeout  = args.telnetWriteTimeout,
		cronTime            = args.cronTime,
		clientType          = args.clientType,
		delayFactor         = args.delayFactor,
		sshMaxLoops         = args.sshMaxLoops,
		jumpHostsFile       = args.jumpHostsFile,
		genMop              = args.genMop,
		strictOrder         = args.strictOrder,
		haltOnError         = args.haltOnError,
		inventoryFile       = args.inventoryFile,
		deviceType          = args.deviceType,
		pluginType          = args.pluginType,
		cmdVerify           = args.cmdVerify,
	)

	### Rady to go ...

	if dictParam['outputJob'] == 0:

		fncRun(dictParam)

	elif (	
		dictParam['outputJob'] == 2 and 
		dictParam['username'] and 
		dictParam['progNumThreads'] and 
		dictParam['logInfo'] and 
		dictParam['useSSHTunnel'] in ['no','yes'] and 
		dictParam['telnetTimeout']
		):

		print("\n#######################################")
		print("# About to run. Ctrl+C if not sure... #")
		print("#######################################\n")
		dictParam['password'] = getpass("### -> PASSWORD (default user: " + dictParam['username'] + "): ")

		fncRun(dictParam)

	else:

		print("Not enough paramteres.\nRun: python script_x_y.py -h for help.\nQuitting...")