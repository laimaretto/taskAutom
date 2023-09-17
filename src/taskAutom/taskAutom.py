#!/usr/bin/env python3

# Copyright (C) 2015-2023 Lucas Aimaretto / laimaretto@gmail.com
#
# This is taskAutom
#
# taskAutom is free software: you can redistribute it and/or modify
# it under the terms of the 3-clause BSD License.
#
# taskAutom is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY of any kind whatsoever.
#

import json
import os
import time
from multiprocessing.pool import ThreadPool
import logging
import importlib
import re
import argparse
from getpass import getpass
import re
import calendar
import sys

# installed
import sshtunnel
import paramiko
from netmiko import ConnectHandler
from netmiko import ConnLogOnly
from scp import SCPClient
import pandas as pd
import yaml
import docx
from docx.enum.style import WD_STYLE_TYPE 
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt


LATEST_VERSION = '8.1.1'

# Constants
IP_LOCALHOST  = "127.0.0.1"
LOG_GLOBAL    = []
LOG_CONSOLE   = []
DICT_PARAM    = dict(
	outputJob        = 0,
	logsDirectory    = None,
	logsCsvFilename  = None,
	logInfo          = None,
	logFileName      = None,
	logsDirTimestamp = None,
	pluginFilename   = None,
	cronTime         = dict(type=None),
	jumpHosts        = dict(),
	pluginType       = 'show',
	cmdVerify        = True,
	auxRetry         = 10,
	inventoryFile    = None,
	useSSHTunnel     = False,
	useHeader        = True,
	strictOrder      = False,
	username         = None,
	password         = None,
	deviceType       = 'nokia_sros',
	readTimeOut      = 10,
	progNumThreads   = 1,
	sshDebug         = False,
	passByRow        = True,
	genMop           = False,
	dataGroupColumn  = 'ip',
	version          = LATEST_VERSION,
	xlsSheetName     = None,
)

# - Parameters per vendor
DICT_VENDOR = dict(
	nokia_sros=dict(
		START_SCRIPT     = "", 
		FIRST_LINE       = "",
		LAST_LINE        = "\nexit all\n",
		FIN_SCRIPT       = "#FINSCRIPT",
		VERSION 	     = "show version", # no \n in the end
		VERSION_REGEX    = "(TiMOS-[A-Z]-\d{1,2}.\d{1,2}.R\d{1,2})",
		HOSTNAME         = "/show chassis | match Name", # no \n in the end
		HOSTNAME_REGEX   = "Name\s+:\s(\S+)",
		HW_TYPE          = "/show chassis | match Type", # no \n in the end
		HW_TYPE_REGEX    = "Type\s+:\s(.+)",
		SHOW             = "",
		SEND_CMD_REGEX   = r"#\s+$",
		MAJOR_ERROR_LIST = ["^FAILED:.+","^ERROR:.+","^Error:.+","invalid token","not allowed"],
		MINOR_ERROR_LIST = ["^MINOR:.+"],
		INFO_ERROR_LIST  = ["^INFO:.+"],
		REMOTE_PORT      = 22,
		SFTP_PORT        = 22,
		SFTP_REGEX_CF    = r"(cf\d+:\/|cf\d+:)",
	),
	md_nokia_sros=dict(
		START_SCRIPT     = "", 
		FIRST_LINE       = "",
		LAST_LINE        = "",
		FIN_SCRIPT       = "",
		VERSION 	     = "show version", # no \n in the end
		VERSION_REGEX    = "(TiMOS-[A-Z]-\d{1,2}.\d{1,2}.R\d{1,2})",
		HOSTNAME         = "/show chassis | match Name", # no \n in the end
		HOSTNAME_REGEX   = "Name\s+:\s(\S+)",
		HW_TYPE          = "/show chassis | match Type", # no \n in the end
		HW_TYPE_REGEX    = "Type\s+:\s(.+)",
		SHOW             = "info json\n",
		SEND_CMD_REGEX   = r"#\s+$",
		MAJOR_ERROR_LIST = ["^FAILED:.+","^ERROR:.+","^Error:.+","invalid token","not allowed"],
		MINOR_ERROR_LIST = ["^MINOR:.+"],
		INFO_ERROR_LIST  = ["^INFO:.+"],
		REMOTE_PORT      = 22,
		SFTP_PORT        = 22,
		SFTP_REGEX_CF    = r"(cf\d+:\/|cf\d+:)",
	),	
	nokia_sros_telnet=dict(
		START_SCRIPT     = "", 
		FIRST_LINE       = "\n/environment no more\n",
		LAST_LINE        = "\nexit all\n",
		FIN_SCRIPT       = "#FINSCRIPT",
		VERSION 	     = "show version", # no \n in the end
		VERSION_REGEX    = "(TiMOS-[A-Z]-\d{1,2}.\d{1,2}.R\d{1,2})",
		HOSTNAME         = "/show chassis | match Name", # no \n in the end
		HOSTNAME_REGEX   = "Name\s+:\s(\S+)",
		HW_TYPE          = "/show chassis | match Type", # no \n in the end
		HW_TYPE_REGEX    = "Type\s+:\s(.+)",
		SHOW             = "",
		SEND_CMD_REGEX   = r"#\s+$",
		MAJOR_ERROR_LIST = ["^FAILED:.+","^ERROR:.+","^Error:.+","invalid token","not allowed"],
		MINOR_ERROR_LIST = ["^MINOR:.+"],
		INFO_ERROR_LIST  = ["^INFO:.+"],
		REMOTE_PORT      = 23,
		SFTP_PORT        = 22,
		SFTP_REGEX_CF    = r"(cf\d+:\/|cf\d+:)",
	),
	nokia_srl=dict(
		START_SCRIPT     = "", 
		FIRST_LINE       = "",
		LAST_LINE        = "",
		FIN_SCRIPT       = "",
		VERSION 	     = "show version", # no \n in the end
		VERSION_REGEX    = "(v\d+.\d+.\d+)",
		HOSTNAME         = "show version", # no \n in the end
		HOSTNAME_REGEX   = "Hostname\s+:\s(\S+)",
		HW_TYPE          = "show version", # no \n in the end
		HW_TYPE_REGEX    = "Chassis Type\s+:\s(.+)",
		SHOW             = "",
		SEND_CMD_REGEX   = None,
		MAJOR_ERROR_LIST = ["^FAILED:.+","^ERROR:.+","^Error:.+","invalid token","not allowed"],
		MINOR_ERROR_LIST = ["^MINOR:.+"],
		INFO_ERROR_LIST  = ["^INFO:.+"],
		REMOTE_PORT      = 22,
		SFTP_PORT        = 22,
		SFTP_REGEX_CF    = r"(cf\d+:\/|cf\d+:)",
	),		
)

####

def fncPrintResults(listOfRouters, timeTotalStart, dictParam):

	separator = "\n------ * ------"

	outTxt    = ""

	outTxt += separator + '\n'

	#### GLOBALS

	outTxt += "Global Parameters:\n"

	outTxt += f'  Template File:              {str(dictParam["pluginFilename"])}\n'
	if bool(dictParam['pluginType']):
		outTxt += f"  Template Type:              {str(dictParam['pluginType'])}\n"
	outTxt += f"  DATA File:                  {str(dictParam['dataFile'])}\n"
	outTxt += f"  DATA UseHeader:             {str(dictParam['useHeader'])}\n"
	outTxt += f"  Folder logInfo:             {dictParam['logInfo']}\n"
	outTxt += f"  Log FileName:               {dictParam['logFileName']}\n"
	outTxt += f"  Text File:                  {dictParam['logInfo']}/job0_{str(dictParam['pluginFileAlone'])}.txt\n"

	if dictParam['genMop'] is True:
		outTxt += f"  MOP filename                {dictParam['logInfo']}/job0_{str(dictParam['pluginFileAlone'])}.docx\n"

	if bool(dictParam['inventoryFile']):
		outTxt += f"  Inventory file              {str(dictParam['inventoryFile'])}\n"


	outTxt += f"  Verify Commands:            {str(dictParam['cmdVerify'])}\n"	
	outTxt += f"  Strict Order:               {str(dictParam['strictOrder'])}\n"
	outTxt += f"  Pass Data By Row:           {str(dictParam['passByRow'])}\n"

	if dictParam['strictOrder'] is True:
		outTxt += f"  Halt-on-Error:              {str(dictParam['haltOnError'])}\n"

	if dictParam['cronTime']['type'] is not None:
		outTxt += f"  CRON Config:                {str(dictParam['cronTime'])}\n"

	if dictParam['strictOrder'] is False:
		outTxt += f"  Total Routers:              {str(len(listOfRouters))}\n"
	else:
		outTxt += f"  Total Lines:                {str(len(listOfRouters))}\n"

	#### CONNECTION

	outTxt += f"\nDefault Connection Parameters:\n"

	if dictParam['inventoryFile'] != None:
		outTxt += f"(Override by inventory file: {dictParam['inventoryFile']})\n\n"
	
	if dictParam['useSSHTunnel'] is True:
		outTxt += f"  Use SSH tunnel:             {str(dictParam['useSSHTunnel'])}; Server: {str(len(dictParam['jumpHosts']))}; File: {dictParam['jumpHostsFile']}\n"
	else:
		outTxt += f"  Use SSH tunnel:             {str(dictParam['useSSHTunnel'])}\n"
	
	outTxt += f"  Total Threads:              {str(dictParam['progNumThreads'])}\n"
	outTxt += f"  Read Timeout:               {str(dictParam['readTimeOut'])}\n"
	outTxt += f"  Time Between Routers:       {str(dictParam['timeBetweenRouters'])}ms\n"
	outTxt += f"  Username:                   {str(dictParam['username'])}\n"
	outTxt += f"  Password Filename:          {str(dictParam['passwordFile'])}\n"
	outTxt += f"  Device Type:                {str(dictParam['deviceType'])}\n"

	if dictParam['outputJob'] > 0:

		timeTotalEnd 	= time.time()

		outTxt += f"{separator}\n"

		df = pd.concat(LOG_GLOBAL)

		timeTotal 		= fncFormatTime(timeTotalEnd - timeTotalStart)
		timeMin         = fncFormatTime(df["time"].min())
		timeAvg         = fncFormatTime(df["time"].mean())
		timeMax         = fncFormatTime(df["time"].max())

		outTxt += f"\nTiming:\n"

		outTxt += f'  timeMin                     {timeMin}\n'
		outTxt += f'  timeAvg:                    {timeAvg}\n'
		outTxt += f'  timeMax:                    {timeMax}\n'
		outTxt += f'  timeTotal:                  {timeTotal}\n'

		outTxt += f"{separator}\n"

		df['threads']     = dictParam['progNumThreads']

		df.to_csv(dictParam['logsCsvFilename'],index=False)

		dfFailed = df[~df['Reason'].isin(['sftpOk','SendSuccess'])]

		if dictParam['strictOrder'] is False:
			outTxt += f"\nFailed routers:             {str(len(dfFailed))}\n"
		else:
			outTxt += f"\nFailed lines:               {str(len(dfFailed))}\n"

		if dictParam['strictOrder'] is True and dictParam['haltOnError'] is True and dictParam['aluLogReason'] not in ['SendSucces','ReadTimeout']:
			outTxt += f"   --> HaltOnError: " + {dictParam['aluLogReason']} + ' <--\n'

		if len(dfFailed) > 0:
			outTxt += dfFailed.to_string(max_colwidth=20)

		outTxt += f"{separator}\n"

		df['Reason'] = df['Reason'].str.replace('(\w+:\d+|\d+.\d+.\d+.\d+:\d{1,6}|\d+.\d+.\d+.\d+)','',regex=True)
		df['Reason'] = df.apply(lambda x: x['Reason'].replace(x['HostName'],''), axis=1)
		dfGroup = df.groupby(['Reason']).agg({'Reason':['count'],'time':['min','max']})

		outTxt += dfGroup.to_string(max_colwidth=20)

		### Final reports
		with open(dictParam['logsDirectory'] + '00_report.txt','w') as f:
			f.write(outTxt)
			f.close()

		with open(dictParam['logsDirectory'] + '00_log_console.txt','w') as f:
			for k in LOG_CONSOLE:
				f.write(k+'\n')
			f.close()

		with open(dictParam['logsDirectory'] + '00_report.json', 'w') as f:
			dictParam['password'] = '*****'
			dictParam.pop('data')
			dictParam.pop('mod')
			dictParam['routersTotal'] = len(df)
			dictParam['routersFailed'] = len(dfFailed)
			if len(dictParam['jumpHosts']) > 0:
				for srv in dictParam['jumpHosts']:
					dictParam['jumpHosts'][srv]['password'] = '*****'
			if len(dictParam['inventory']) > 0:
				for ip in dictParam['inventory']:
					dictParam['inventory'][ip]['password'] = '*****'
			dictParam['timing'] = {}
			dictParam['timing']['min'] = timeMin
			dictParam['timing']['avg'] = timeAvg
			dictParam['timing']['max'] = timeMax
			dictParam['timing']['total'] = timeTotal
			json.dump(dictParam, f)
			f.close()

		### Data File of failed routers
		dataFile = pd.read_excel(dictParam['dataFile'],sheet_name=dictParam['xlsSheetName']) if dictParam['xlsSheetName'] is not None else pd.read_csv(dictParam['dataFile'])
		grpCol   = dictParam['dataGroupColumn']
		listFailed = dfFailed['IP'].to_list()
		failedDataFile = dataFile[dataFile[grpCol].isin(listFailed)]
		if len(failedDataFile)>0:
			failedDataFile.to_csv(dictParam['logsDirectory'] + '00_failedDataFile.csv',index=False)

	outTxt += f"{separator}\n"

	print(outTxt)

def fncFormatTime(timeFloat, adjust=True):

	move = 100

	if adjust==True:

		unit = 's'

		if timeFloat > 120:
			timeFloat = timeFloat / 60
			unit = 'm'

		timeFloat = float(int(timeFloat*move))/move	

		return str( timeFloat ) + unit

	else:

		return float(int(timeFloat*move))/move

def fncPrintConsole(inText, show=1):

	localtime   = time.localtime()
	if show:
		output = str(time.strftime("%H:%M:%S", localtime)) + "| " + inText
		print(output)
		LOG_CONSOLE.append(output)

def getListOfRouters(dictParam):
	"""
	Function to obtain the unique list of routers.
	If using headers, the data will be otained from such column
	If not using headers, data will be otained from the first column.

	## If using strictOrder, we get the list of routers as is, with out any filtering ##
	## This is so, because later on, we do a for-loop on this list.

	Args:
		dictParam: configuration parameters of taskAutom

	Returns:
		[list]: list of routers.
	"""

	groupColumn = dictParam['dataGroupColumn']
	data        = dictParam['data']
	outputJob   = dictParam['outputJob']

	dOut = {}

	if outputJob in [0,2]:
		# If strictOrder is True, we get the list of routers
		# as it is defined inside the dataFile
		if dictParam['strictOrder'] is True:

			if dictParam['useHeader'] is True:
				try:
					routers = list(data[groupColumn])
				except Exception as e:
					print("No column header " + str(e) + " in file " + dictParam['dataFile'] + ". Quitting...\n")
					quit()
			else:
				routers = list(data[0])

		else:
		# If strictOrder is False, we get the unique list of routers
		# as it is defined inside the dataFile. The order here is not important.

			if dictParam['useHeader'] is True:
				try:
					routers = list(data[groupColumn].unique())
				except Exception as e:
					print("No column header " + str(e) + " in file " + dictParam['dataFile'] + ". Quitting...\n")
					quit()				
			else:
				routers = list(data[0].unique())

	else:

		routers = []

		try:
			for row in data.itertuples():
				routers.append((row.ip,row.ftpLocalFilename,row.ftpRemoteFilename))
		except Exception as e:
			print("Something happened with the data file " + dictParam['dataFile'] + ".\n" + str(e) + ".\nQuitting...")
			quit()		

	# We build a dictionary with per-router connection info.
	for info in routers:

		if outputJob in [0,2]:
			router = info
		else:
			router = info[0]
			fLocal = info[1]
			fRemot = info[2]

		if router not in dOut.keys():

			dOut[router] = {
				'username':dictParam['username'],
				'password':dictParam['password'],
				'deviceType':dictParam['deviceType'],
				'useSSHTunnel':dictParam['useSSHTunnel'],
				'readTimeOut':dictParam['readTimeOut'],
				'jumpHost':None,
				'systemIP':router,
			}

		if outputJob in [0,2]:
			dOut[router]['pluginScript'] = []
		else:
			try:
				dOut[router]['ftpFiles'].append((fLocal,fRemot))
			except:
				dOut[router]['ftpFiles'] = [(fLocal,fRemot)]

	return routers, dOut

def verifyCronTime(cronTime):
	"""[We verify cronTime before moving on]

	Args:
		cronTime ([list]): [list of parameters]

	Returns:
		[list]
	"""

	dCron = {}

	if cronTime in ['',[],None]:
		dCron['type'] = None
		return dCron
	else:
		dCron['type'] = cronTime[0]
		if dCron['type'] not in ['oneshot','periodic']:
			print("CronType con only be either 'oneshot' or 'periodic'. Quitting ...")
			quit()
		
		if dCron['type'] == 'oneshot':
			if len(cronTime)!=7:
				print('Wrong cronTime length for "oneshot". Quitting ...')
				quit()
			else:
				dCron['cronName']   = str(cronTime[1])
				dCron['month']      = str(cronTime[2])
				dCron['weekday']    = str(cronTime[3])
				dCron['dayOfMonth'] = int(cronTime[4])
				dCron['hour']       = int(cronTime[5])
				dCron['minute']     = int(cronTime[6])

		elif dCron['type'] == 'periodic':
			if len(cronTime)!=3:
				print('Wrong cronTime length for "periodic". Quitting ...')
				quit()
			else:
				dCron['cronName']   = str(cronTime[1])
				dCron['interval']   = int(cronTime[2])			

	if dCron['cronName'][0] in [str(x) for x in range(0,10)]:
		print('Wrong CRON name. First char cannot be a number. Quitting ...')
		quit()		
	elif not re.compile(r'^\S+$').search(dCron['cronName']):
		print('Wrong CRON name. Quitting ...')
		quit()
	
	if dCron['type'] == 'oneshot':

		if dCron['month'] not in [calendar.month_name[x].lower() for x in range(1,13)]:
			print('Wrong month name. Quitting ...')
			quit()

		if dCron['weekday'] not in [calendar.day_name[x].lower() for x in range(0,7)]:
			print('Wrong weekDay name. Quitting ...')
			quit()		

		if dCron['dayOfMonth'] not in list(range(1,32)):
			print('Wrong dayOfMonth value. Quitting ...')
			quit()			

		if dCron['hour'] not in list(range(0,24)):
			print('Wrong hour value. Quitting ...')
			quit()

		if dCron['minute'] not in list(range(0,60)):
			print('Wrong minute value. Quitting ...')
			quit()

	elif dCron['type'] == 'periodic':

		if dCron['interval'] not in list(range(30,42949672)):
			print('Wrong interval value. Quitting ...')
			quit()

	return dCron

def verifyServers(jumpHostsFile):
	"""We verify the SERVERS dictionary before moving on.

	Args:
		SERVERS ([str]): [Name of the file containing servers information in YML format.]

	Returns:
		[dict]: [Dictionary with servers information]
	"""

	try:
		with open(jumpHostsFile,'r') as f:
			servers = yaml.load(f, Loader=yaml.FullLoader)
			f.close()
	except:
		print("Missing " + jumpHostsFile + " file. Quitting..")
		quit()

	fields = ['name','user','password','ip','port']

	for k in servers.keys():
		for f in fields:
			if f in servers[k].keys():
				if not servers[k][f]:
					print('Missing value for field "' + str(f) + '" in server "' + str(k) + '". Quitting...')
					quit()
			else:
				print('Missing field "' + str(f) + '" in server "' + str(k) + '". Quitting...')
				quit()

	# If before checking is ok, we create a new dictionary with correlative keys for those values...
	return servers

def verifyData(dictParam):
	"""[Verify DATA file]

	Args:
		csvFile ([str]): [Name of CSV file]

	Returns:
		[list]: [List of Routers]
	"""

	if dictParam['useHeader'] is True:
		useHeader = 0
	else:
		useHeader = None

	if dictParam['xlsSheetName'] is None:
	
		# We have CSV
		try:
			routers = pd.read_csv(dictParam['dataFile'], header=useHeader)
		except Exception as e:
			print(e)
			print("Error trying to open file " + dictParam['dataFile'] + ". Quitting...\n")
			quit()
	
	else:

		# We have XLSX
		try:
			routers = pd.read_excel(dictParam['dataFile'], sheet_name=dictParam['xlsSheetName'], header=useHeader)
		except Exception as e:
			print(e)
			print("Error trying to open file " + dictParam['dataFile'] + ". Quitting...\n")
			quit()

	return routers

def verifyPlugin(pluginFilename):
	"""[Verifies the plugin template]

	Args:
		pluginFilename ([str]): [Name of config template]

	Returns:
		[module]: [The module]
	"""

	try:
		if pluginFilename.split(".")[-1] == "py":
			# if '/' in pluginFilename:
			# 	pluginFilename = pluginFilename.replace('/','.')
			print(pluginFilename)
			spec = importlib.util.spec_from_file_location("construir_cliLine",pluginFilename)
			mod  = importlib.util.module_from_spec(spec)
			sys.modules["construir_cliLine"] = mod
			spec.loader.exec_module(mod)
			#mod  = importlib.import_module(".","showInterface")
			print(mod)
		else:
			print("Missing config file. Verify extension of the file to be '.py'. Quitting...")
			quit()
	except Exception as e:
		print(e)
		print("----\nError importing plugin. Quitting ...")
		quit()

	return mod

def verifyConfigFile(config_file):
	""" This function checks the whole text in order to search for ASCII 
	characters (7bit) since 8bit chars won't allow a proper boot process.
	"""

	charset_allowed = [chr(c) for c in range(128)]

	for i,line in enumerate(config_file.split('\n')):
		for character in line:
			if character not in charset_allowed:
				return i+1, line, character

	return -1,-1

def verifyInventory(dictParam):

	inventoryFile = dictParam['inventoryFile']
	jumpHostsFile = dictParam['jumpHostsFile']
	outputJob     = dictParam['outputJob']

	if outputJob in [0,2]:
		listOfRouters = dictParam['listOfRouters']
	else:
		listOfRouters = [x[0] for x in dictParam['listOfRouters']]

	columns = ['ip','username','password','deviceType','useSSHTunnel','readTimeOut','jumpHost']

	try:
		df = pd.read_csv(inventoryFile)
	except:
		print("Inventory: The file " + inventoryFile + " was not found. Quitting ...")
		quit()

	for col in df.columns:
		if col not in columns:
			print("Inventory: The field '" + col +  "' is not a valid one. Valids are: " + str(columns) + ". Quitting...")
			quit()

	for col in columns:
		if col not in df.columns:
			print("Inventory: The inventory file ("+inventoryFile+") is missing the field '" + col +  "'. Quitting...")
			quit()			

	df2 = df.copy()
	df2 = df2.fillna("")

	dOut   = {}

	df2 = df2[df2.ip.isin(listOfRouters)]

	for row in df2.itertuples():

		ip       = row.ip
		jh       = row.jumpHost if row.jumpHost != '' else None
		dt       = row.deviceType if row.deviceType != '' else dictParam['deviceType']
		tun      = row.useSSHTunnel if row.useSSHTunnel != '' else dictParam['useSSHTunnel']
		rto      = row.readTimeOut if row.readTimeOut != '' else dictParam['readTimeOut']
		username = row.username if row.username != '' else dictParam['username']
		password = row.password if row.password != '' else dictParam['password']

		if tun not in ['yes','no','',None]:
			print(f'Inventory: The router {ip} is not using a valid sshTunnel option. For default, leave empty. Quitting...')
			quit()
		else:
			tun = True if row.useSSHTunnel == 'yes' else False

		if tun is True:

			serversList = list(verifyServers(jumpHostsFile).keys())

			if jh not in serversList+ ['',None]:
				print(f'Inventory: The router {ip} is using sshtunnel and has not a valid jumpHost.\nIf empty, using default. Available servers inside the file {jumpHostsFile}: {str(serversList)}.\nQuitting...')
				quit()

		if dt not in list(DICT_VENDOR.keys()) + ['',None]:
			print(f'Inventory: The router {ip} is not using a valid deviceType. For default, leave empty. Quitting...')
			quit()

		if rto != '':
			try:
				int(rto)
			except:
				print(f'Inventory: The router {ip} has not a valid ReadTimeOut. For default, leave empty. Quitting...')
				quit()

		dOut[ip] = {
			'username':username,
			'password':password,
			'deviceType':dt,
			'useSSHTunnel':tun,
			'readTimeOut':rto,
			'jumpHost':jh,
			'systemIP':ip,
		}

	return dOut

def renderMop(aluCliLineJob0, dictParam):
	"""[Generates a MOP based on the CSV and plugin information]

	Args:
		aluCliLineJob0 ([file]): [configLines]
		pluginFilename ([str]):  [The plugin for this MOP]

	Returns:
		None
	"""

	# Verify if logsDirectory exists.
	if not os.path.exists(dictParam['logInfo']):
		os.makedirs(dictParam['logInfo'])

	job0docx = dictParam['logInfo'] + "/job0_" + dictParam['pluginFileAlone'] + ".docx"
	job0text = dictParam['logInfo'] + "/job0_" + dictParam['pluginFileAlone'] + ".txt"

	if dictParam['genMop'] is True:

		print("\nGenerating MOP: " + job0docx)
		config = aluCliLineJob0.split('\n')
		config = [x for x in config if len(x) > 0]

		myDoc = docx.Document()
		myStyles = myDoc.styles  

		styleConsole = myStyles.add_style('Console', WD_STYLE_TYPE.PARAGRAPH)
		styleConsole.font.name = 'Courier'
		styleConsole.font.size = Pt(9)
		styleConsole.paragraph_format.keep_together = True

		styleConsole.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		#styleConsole.paragraph_format.line_spacing = Pt(10)
		#styleConsole.paragraph_format.line_spacing = .2
		styleConsole.paragraph_format.space_after = Pt(2)

		myDoc.add_heading('MOP for ' + dictParam['pluginFilename'], 0)

		for i,row in enumerate(config):

			if i == 0:
				myDoc.add_heading('Configuraciones',1)

			if 'Heading_2' in row.split(":")[0]:
				row = ''.join(row.split(":")[1:])
				subtitle = myDoc.add_paragraph(row)
				subtitle.style = myDoc.styles['Heading 2']
				subtitle.paragraph_format.line_spacing = 1.5

			elif 'Heading_3' in row.split(":")[0]:
				row = ''.join(row.split(":")[1:])
				subtitle = myDoc.add_paragraph(row)
				subtitle.style = myDoc.styles['Heading 3']
				subtitle.paragraph_format.line_spacing = 1.5

			else:
				configText = myDoc.add_paragraph(row)
				configText.style = myDoc.styles['Console']

		myDoc.save(job0docx)
		print("MOP done...")

	with open(job0text,'w') as f:
		f.write(aluCliLineJob0)
		f.close()

def renderCliLine(IPconnect, dictParam, i):
	"""
	This function renders the script, based both on the Data file (data) and the plugin (py)
	There are several possibilities for treating the data, depending on the following.

	- groupColumn: how data is filtered
	- strictOrder: if we follow the order described inside the data file
	- useHeader: do we pay attention the headers or use numerals instead?
	- passByRow: do we pass to the module the complete data or only rowByRow? Only valid when strictOrder=False.

	Args:
		IPconnect (_type_): IP of router (or eventually, the grouped data value)
		dictParam (_type_): dictionary of parameters
		mod (_type_):       plugin
		data (_type_):      Pandas DataFrame
		i (_type_):         id which identifies the IPconnect value

	Returns:
		_type_: _description_
	"""

	#   i             0         1        2        3             4                 5     6  7
	#   0     10.3.0.41    ZONA_X  0.0.0.0  0.0.1.3      ROUTER_A  TiMOS-C-16.0.R6      1  4
	#   1     10.3.0.42    ZONA_Y  0.0.0.0  0.0.1.3      ROUTER_B  TiMOS-C-16.0.R6   9886  4
	#   2     10.3.0.43    ZONA_Y  0.0.0.0  0.0.1.3      ROUTER_C  TiMOS-C-16.0.R6   9886  4
	#   3     10.3.0.44    ZONA_Y  0.0.1.3  0.0.1.3      ROUTER_D  TiMOS-B-7.0.R7    9886  4
	#   4     10.3.0.45    ZONA_Y  0.0.1.3  0.0.1.3      ROUTER_E  TiMOS-B-7.0.R7    9886  4
	#   5     10.3.0.46    ZONA_Y  0.0.1.3  0.0.1.3      ROUTER_F  TiMOS-B-9.0.R3    9886  4
	#   6     10.3.0.47    ZONA_Y  0.0.1.3  0.0.1.3      ROUTER_G  TiMOS-B-9.0.R3    9886  4
	#   7     10.3.0.48    ZONA_Y  0.0.1.3  0.0.1.3      ROUTER_H  TiMOS-B-7.0.R7    9886  4
	#   8     10.3.0.49    ZONA_Y  0.0.1.3  0.0.1.3      ROUTER_I  TiMOS-C-19.10.R8  9886  4
	#   9     10.3.0.50    ZONA_Y  0.0.1.3  0.0.1.3      ROUTER_J  TiMOS-B-7.0.R7    9886  4


	aluCliLine     = ""
	groupColumn    = dictParam['dataGroupColumn']
	jobType        = dictParam['outputJob'] 
	strictOrder    = dictParam['strictOrder']
	useHeader      = dictParam['useHeader']
	pluginFilename = dictParam['pluginFilename']
	dataFile       = dictParam['dataFile']
	passByRow      = dictParam['passByRow']
	mod            = dictParam['mod']
	data           = dictParam['data']

	if jobType == 2:
		mop = None
	elif jobType == 0:
		if i == -1:
			mop = None
		else:
			mop = 1

	if strictOrder is False:

		# Since strictOrder = no, then we pass to the module
		# all the data, filterd by IPconnect

		if useHeader is True:
			pluginData = data[data[groupColumn] == IPconnect]
		else:
			pluginData = data[data[0] == IPconnect]

		if passByRow is True:

			# since passByRow is True, we pass the data to the module row by row
			# hence we apply in here the itertuples() for-loop, which is the default mode.
			# The length of data is len(data)
			# The row order is 'j'

			for j, row in enumerate(pluginData.itertuples()):
				try:
					aluCliLine = aluCliLine + mod.construir_cliLine(j, row, len(pluginData), mop)
				except Exception as e:
					print(f'\nError trying to use plugin {pluginFilename}.\nVerify variables inside of it, or the data file {dataFile}.')
					print('\nError: ' + str(e))
					print('Row: ' + str(row))		
					print('Quitting...')
					sys.exit(1)

		else:

			# since passByRow is False, we pass the complete.
			# the itertuples() for-loop, will be needed inside the plugin.
			# this is only possible with strictMode == no.
			# The length of data is len(data)
			# The row order is 'j'
			try:
				aluCliLine = aluCliLine + mod.construir_cliLine(i, pluginData, len(pluginData), mop)
			except Exception as e:
				print(f'\nError trying to use plugin {pluginFilename}.\nVerify variables inside of it, or the data structure file, {dataFile}, looking for inconsistencies.')
				print('\nError: ' + str(e))				
				sys.exit(1)			

	else:

		# Since strictOrder = True, then we pass to the module
		# all the data, row by row, by id i, which comes from 
		# fncRun(). 
		# Then the length of data is always 1.
		# The row order is 0

		try:
			pluginData = list(data.itertuples())[i]
			aluCliLine = mod.construir_cliLine(0, pluginData, 1, mop)
		except Exception as e:
			print(f'\nError trying to use plugin {pluginFilename}.\nVerify variables inside of it, or the data file {dataFile}.')
			print('\nError: ' + str(e))
			print('Row: ' + str(pluginData))			
			print('Quitting...')
			sys.exit(1)

	return aluCliLine


def buildScripts(dictParam):
	"""
	This function builds scripts per router. This function must be used
	when using taskAutom as a library (import)

	Args:
		dictParam

	Returns:
		_dict with results
	"""

	dictParam['data'] = verifyData(dictParam)
	dictParam['mod']  = verifyPlugin(dictParam['pluginFilename'])
	dictParam['pluginFileAlone'] = dictParam['pluginFilename'].split('/')[-1]
	dictParam['listOfRouters'], _ = getListOfRouters(dictParam)

	d = {}

	for IPconnect in dictParam['listOfRouters']:
		if IPconnect not in d.keys():
			d[IPconnect] = {}
		d[IPconnect][dictParam['pluginFileAlone']] = renderCliLine(IPconnect, dictParam, 1)

	return d
###

def run_mi_thread(i, routerInfo, dictParam):
	"""[summary]

	Args:
		i ([type]): [description]
		ip ([type]): [description]
		dictParam ([dict]): [Dictionary with connection parameters]
		pluginScript ([type]): [description]
		ftpFileName
	"""

	outRx = myConnection(i, routerInfo, dictParam).run()

	return outRx

class myConnection():
	"""
	[Class for connection Object]
	"""

	def __init__(self, thrdNum, routerInfo, dictParam):

		self.outputJob 	      = dictParam['outputJob']
		self.logsDirectory    = dictParam['logsDirectory']
		self.logsCsvFilename  = dictParam['logsCsvFilename']
		self.logInfo          = dictParam['logInfo']
		self.logFileName      = dictParam['logFileName']
		self.logsDirTimestamp = dictParam['logsDirTimestamp']
		self.plugin           = dictParam['pluginFilename']

		# local generated variables
		self.connInfo = {
			'localPort':-1,
			'remotePort':-1,
			'controlPlaneAccess': False,
			'aluLogged': False,
			'aluLogReason':"N/A",
			'hostname':"N/A",
			'timos':"N/A",
			'hwType':"N/A",
			'cronTime':dictParam['cronTime'],
			'sshServer': None,
			'conn2rtr': None,
			'jumpHosts':dictParam['jumpHosts'],
			'pluginType':dictParam['pluginType'],
			'cmdVerify':dictParam['cmdVerify'],
			'tDiff':0,
			'runStatus':-1, # revisar, solo una vez
			'strConn': "Con-" + str(thrdNum) + "| ",
			'num':thrdNum,
			'outRx':'',
			'outRxJson':{},
			'cronScript':None,
			'auxRetry':dictParam['auxRetry'],
			'sshDebug':dictParam['sshDebug']
		}

		self.connInfo.update(routerInfo)

		# We update the outputjob relevant information...
		if self.outputJob == 2:
			self.connInfo['pluginScript'] = DICT_VENDOR[self.connInfo['deviceType']]['START_SCRIPT'] + \
											DICT_VENDOR[self.connInfo['deviceType']]['FIRST_LINE'] + \
											routerInfo['pluginScript'][-1] + \
											DICT_VENDOR[self.connInfo['deviceType']]['LAST_LINE'] + \
											DICT_VENDOR[self.connInfo['deviceType']]['FIN_SCRIPT']				
		elif self.outputJob == 3:
			self.connInfo['ftpFiles'] = routerInfo['ftpFiles']
			self.connInfo['ftpTotalTxFiles'] = 0

		# Do we you use jumpHosts?
		if self.connInfo['useSSHTunnel'] is True or dictParam['inventoryFile'] != None:
			self.connInfo['jumpHost'] = [x for i,x in enumerate(self.connInfo['jumpHosts']) if self.connInfo['num'] % len(self.connInfo['jumpHosts']) == i][0]
		else:
			self.connInfo['jumpHost'] = -1

		# SFTP Port
		self.connInfo['sftpPort'] = DICT_VENDOR[self.connInfo['deviceType']]['SFTP_PORT']

		# Identify connection ports
		if ":" in self.connInfo['systemIP']:
			self.connInfo['remotePort'] = int( self.connInfo['systemIP'].split(":")[1] )			
			self.connInfo['systemIP']   = self.connInfo['systemIP'].split(":")[0]
		else:
			self.connInfo['remotePort'] = DICT_VENDOR[self.connInfo['deviceType']]['REMOTE_PORT']

		# --- Users
		self.ROUTER_USER1    = [self.connInfo['username'],self.connInfo['password']]
		self.ROUTER_USER2    = ["extraUser1","extraPassword1"]
		self.ROUTER_USER3    = ["extraUser2","extraPassword2"]
		self.ROUTER_USER     = [self.ROUTER_USER1]

	def run(self):

		# We update the connection info dictionary, after we've set up the connection towards the router...
		self.connInfo = self.fncConnectToRouter(self.connInfo)

		if bool(self.connInfo['conn2rtr']) is True and self.connInfo['aluLogged'] is True:
			
			fncPrintConsole(self.connInfo['strConn'] + "#### Auth ok for " + self.connInfo['systemIP'] +  " ...")

			self.connInfo['timos']      = self.fncAuxGetVal(self.connInfo, 'timos')
			self.connInfo['hostname']   = self.fncAuxGetVal(self.connInfo, 'hostname')
			self.connInfo['timosMajor'] = self.fncAuxGetVal(self.connInfo, 'timosMajor')
			self.connInfo['hwType']     = self.fncAuxGetVal(self.connInfo, 'hwType')
			
			if self.outputJob == 2:

				fncPrintConsole(self.connInfo['strConn'] + "#### Running routine for " + self.connInfo['systemIP'] +  " ...")

				if self.connInfo['cronTime']['type'] is not None:

					self.connInfo = self.fncUploadFile(self.connInfo)

					if self.connInfo['sftpStatus'] is True:

						self.connInfo = self.runCron(self.connInfo)
						self.connInfo = self.routerRunRoutine(self.connInfo)

				else:
					
					self.connInfo = self.routerRunRoutine(self.connInfo)		

			elif self.outputJob == 3:

				self.connInfo = self.fncUploadFile(self.connInfo)				

			fncPrintConsole(self.connInfo['strConn'] + "End-of-run: " + str(self.connInfo['aluLogReason']))

		self.connInfo = self.logData(self.connInfo, self.logInfo, self.logsDirTimestamp, self.plugin, self.logsDirectory)

		#######################
		# closing connections #

		if bool(self.connInfo['conn2rtr']) is True or self.connInfo['aluLogged'] is True:
			self.connInfo['conn2rtr'].disconnect()

		if self.connInfo['useSSHTunnel'] is True and bool(self.connInfo['sshServer']) is True:
			self.connInfo['sshServer'].stop(force=True)

		#                     #
		#######################

		return self.connInfo

	def fncWriteToConnection(self, inText, connInfo):

		conn2rtr           = connInfo['conn2rtr']
		pluginType         = connInfo['pluginType']
		readTimeOut        = connInfo['readTimeOut']
		cmdVerify          = connInfo['cmdVerify']
		deviceType         = connInfo['deviceType']

		expectString       = DICT_VENDOR[connInfo['deviceType']]['SEND_CMD_REGEX']

		outputTxt  = ''
		outputJson = {}

		mdDevice           = re.match('^md_',deviceType)
		mdShow             = DICT_VENDOR[connInfo['deviceType']]['SHOW']


		# ### Writes to a connection. 
		if isinstance(inText,list):

			if pluginType == 'config':

				try:
					outputTxt    = conn2rtr.send_config_set(config_commands=inText, enter_config_mode=False, cmd_verify=cmdVerify, read_timeout=readTimeOut)
					aluLogReason = ""
					runStatus    = 1
				except Exception as e:
					aluLogReason = str(e).replace('\n',' ').lstrip()
					runStatus    = -1						

			elif pluginType == 'show':

				try:

					for cmd in inText:
						if not mdDevice:
							rx        = conn2rtr.send_command(cmd, expect_string=expectString, cmd_verify=cmdVerify, read_timeout=readTimeOut)
							outputTxt = outputTxt + '\n' + cmd + '\n' + rx
							outputJson[cmd] = rx
						else:
							_        = conn2rtr.send_command(cmd, expect_string=expectString, cmd_verify=cmdVerify, read_timeout=readTimeOut)
							rx       = conn2rtr.send_command(mdShow, expect_string=expectString, cmd_verify=cmdVerify, read_timeout=readTimeOut)
							d = json.loads(rx)
							outputJson[cmd] = d
							outputTxt = outputTxt + '\n' + cmd + '\n' + rx				
					
					aluLogReason = ""
					runStatus    = 1

				except Exception as e:
					outputTxt = outputTxt + '\n' + cmd + '\n' + rx
					aluLogReason = str(e).replace('\n',' ').lstrip()
					runStatus    = -1

		elif isinstance(inText,str):
			
			try:
				outputTxt    = conn2rtr.send_command(inText, expect_string=expectString, cmd_verify=cmdVerify, read_timeout=readTimeOut)
				aluLogReason = ""
				runStatus    = 1					
			except Exception as e:
				outputTxt    = ''
				aluLogReason = str(e).replace('\n',' ').lstrip()
				runStatus    = -1

		return runStatus, aluLogReason, outputTxt, outputJson

	def fncAuxGetVal(self, connInfo, what):

		def _getData(inText,connInfo):

			auxRetry = connInfo['auxRetry']

			for i in range(auxRetry):
				runStatus, aluLogReason, rx, _ = self.fncWriteToConnection(inText, connInfo)
				if runStatus == 1:
					break

			return rx

		if what == "timos":

			inText  = DICT_VENDOR[connInfo['deviceType']]['VERSION']
			inRegex = DICT_VENDOR[connInfo['deviceType']]['VERSION_REGEX']
			rx      = _getData(inText,connInfo)
			match   = re.compile(inRegex).search(rx)

			try:
				timos   = match.groups()[0]
			except:
				timos   = "not-matched"

			return timos

		elif what == 'hostname':

			inText  = DICT_VENDOR[connInfo['deviceType']]['HOSTNAME']
			inRegex = DICT_VENDOR[connInfo['deviceType']]['HOSTNAME_REGEX']
			rx      = _getData(inText,connInfo)			
			match   = re.compile(inRegex).search(rx)

			try:
				hostname = match.groups()[0]
			except:
				hostname = "host_" + str(connInfo['num']) + "_not-matched"

			return hostname

		elif what == 'hwType':

			inText  = DICT_VENDOR[connInfo['deviceType']]['HW_TYPE']
			inRegex = DICT_VENDOR[connInfo['deviceType']]['HW_TYPE_REGEX']
			rx      = _getData(inText,connInfo)			
			match   = re.compile(inRegex).search(rx)

			try:
				hwType = match.groups()[0]
			except:
				hwType = "not-matched"

			return hwType			

		elif what == "timosMajor":

			try:
				timosMajor = int(self.connInfo['timos'].split("-")[2].split(".")[0])
			except:
				timosMajor = "not-matched"

			return timosMajor	

	def fncConnectToRouter(self, connInfo):
		"""[We update the connection info dictionary, after we've set up the connection towards the router]

		Args:
			connInfo ([dict]): [Contains all conection related relevant information ]

		Returns:
			[dict]: [Updated connInfo dictionary]
		"""

		### SSH Tunnel

		if connInfo['useSSHTunnel'] is True:

			connInfo = self.fncSshServer(connInfo)

		else:

			fncPrintConsole(connInfo['strConn'] + "Using direct " + connInfo['deviceType'] + " access: ")
			fncPrintConsole(connInfo['strConn'] + "Trying router " + connInfo['systemIP'] + ":" + str(connInfo['remotePort']) )

			connInfo['controlPlaneAccess'] 	= True	
			connInfo['localPort'] 			= connInfo['remotePort']
			connInfo['sshServer']    		= None

		### Connect to router

		if connInfo['controlPlaneAccess'] is True:

			connInfo = self.routerLogin(connInfo)

		else:

			connInfo['conn2rtr']     = None
			connInfo['aluLogged'] 	 = False
			connInfo['username']     = "N/A"
			connInfo['password']     = "N/A"

		return connInfo

	def fncUploadFile(self, connInfo):
		### upload configFile via SFTP

		def setTransport(ip, sftpPort):

			transport = paramiko.Transport((ip,sftpPort))
			transport.connect(None,connInfo['username'],connInfo['password'])
			
			# The routers with timos above 6.X do support SFTP.
			# Otherwise we need to use SCP.
			if connInfo['timosMajor'] != 'not-matched':
				if connInfo['timosMajor'] > 6:
					fncPrintConsole(connInfo['strConn'] + "uploading via SFTP: " + str(sftpPort))
					sftp = paramiko.SFTPClient.from_transport(transport)
				else:
					fncPrintConsole(connInfo['strConn'] + "uploading via SCP: " + str(sftpPort))
					sftp = SCPClient(transport)
			else:
				fncPrintConsole(connInfo['strConn'] + "TiMOS not-matched. Asuming SCP. Uploading via SCP: " + str(sftpPort))
				sftp = SCPClient(transport)

			return transport, sftp

		def sendFiles(sftp,ftpFiles):

			for i, (fileLocal,fileRemote) in enumerate(ftpFiles):

				match = re.match(DICT_VENDOR[connInfo['deviceType']]['SFTP_REGEX_CF'], fileRemote)

				if not match:
					fileRemote = "cf3:/" + fileRemote
					fncPrintConsole(connInfo['strConn'] + "no CF specified; assuming cf3. Uploading file: " + fileLocal + "->" + fileRemote)
				else:
					fncPrintConsole(connInfo['strConn'] + "Uploading file: " + fileLocal + "->" + fileRemote)

				try:
					sftp.put(fileLocal,fileRemote)
					sftpStatus   = True
					aluLogReason = 'sftpOk'
				except Exception as e:
					print(str(e))
					sftpStatus   = False
					aluLogReason = str(e)
					return sftpStatus, aluLogReason, fileRemote, i

				i = i + 1

			return sftpStatus, aluLogReason, fileRemote, i

		if self.outputJob == 2:

			datos      = connInfo['pluginScript']
			fileRemote = connInfo['hostname'] + "_commands.cfg"
			fileLocal  = self.logsDirectory + fileRemote

			ftpFiles   = [(fileLocal,fileRemote)]

			# We write here the contents of the data to be run inside the CRON
			# We hence don't log it thereafter.
			with open(fileLocal,'w') as fc:
				fc.write(datos)
				fc.close()

		elif self.outputJob == 3:

			ftpFiles  = connInfo['ftpFiles']

		## Setting up the connection ...

		if connInfo['useSSHTunnel'] is True:

			sshSftp       = self.fncSshServer(connInfo, sftp=True)
			sftpPort      = sshSftp['localPort']
			sshServerSftp = sshSftp['sshServer']

			transport, sftp = setTransport(IP_LOCALHOST, sftpPort)
			sftpStatus, aluLogReason, fileRemote, i = sendFiles(sftp,ftpFiles)

			sftp.close()
			transport.close()
			sshServerSftp.stop()			

		else:
			sftpPort = connInfo['sftpPort']
			transport, sftp = setTransport(connInfo['systemIP'], sftpPort)
			sftpStatus, aluLogReason, fileRemote, i = sendFiles(sftp,ftpFiles)

			sftp.close()
			transport.close()

		connInfo['sftpStatus']      = sftpStatus
		connInfo['aluLogReason']    = aluLogReason
		connInfo['ftpRemoteFile']   = fileRemote
		connInfo['ftpTotalTxFiles'] = i

		return connInfo

	def fncSshServer(self, connInfo, sftp=False):

		controlPlaneAccess = False
		localPort 		   = None
		server             = None
		aluLogReason       = 'no-ssh-tunnel'		

		jumpHost = connInfo['jumpHost']
		servers  = connInfo['jumpHosts']

		tempIp   = servers[jumpHost]['ip']
		tempPort = servers[jumpHost]['port']
		tempUser = servers[jumpHost]['user']
		tempPass = servers[jumpHost]['password']

		if sftp:
			remotePort = connInfo['sftpPort']
		else:
			remotePort = connInfo['remotePort']

		systemIP = connInfo['systemIP']
		fncPrintConsole(connInfo['strConn'] + "Trying sshServer on IP: " + str(tempIp))
		
		try:
			server = sshtunnel.SSHTunnelForwarder( 	(tempIp, tempPort), 
												ssh_username = tempUser, 
												ssh_password = tempPass, 
												remote_bind_address = (systemIP, remotePort),
												allow_agent = False,
											)
		except Exception as e:
			aluLogReason = "Problems creating SSH server: " + str(e).replace('\n',' ').lstrip()
			fncPrintConsole(connInfo['strConn'] + str(aluLogReason))
			server.stop(force=True)
			controlPlaneAccess = False
			localPort          = None
			server             = None

		if server is not None:

			try:
				server.start()
				localPort = server.local_bind_port
				fncPrintConsole(connInfo['strConn'] + "Trying sshServerTunnel on port: " + str(localPort))		
			except Exception as e:
				fncPrintConsole(connInfo['strConn'] + "Trying sshServerTunnel on port: " + str(localPort))
				aluLogReason = "Problems starting SSH server: " + str(e).replace('\n',' ').lstrip()
				fncPrintConsole(connInfo['strConn'] + aluLogReason)
				server.stop(force=True)
				controlPlaneAccess = False
				localPort          = None
				server             = None
		
		if server is not None:

			fncPrintConsole(connInfo['strConn'] + "Trying router " + IP_LOCALHOST + ":" + str(localPort) + " -> " + connInfo['systemIP'] + ":" + str(connInfo['remotePort']))				

			server.check_tunnels()

			if server.tunnel_is_up[('0.0.0.0',localPort)] is False:
				aluLogReason = 'SSH Error: Tunnel is not up.'
				fncPrintConsole(connInfo['strConn'] + aluLogReason)
				server.stop(force=True)
				controlPlaneAccess = False
				localPort          = None
				server             = None
			else:
				controlPlaneAccess = True

		connInfo['aluLogReason']       = aluLogReason
		connInfo['controlPlaneAccess'] = controlPlaneAccess
		connInfo['localPort']          = localPort
		connInfo['sshServer']          = server

		return connInfo

	def routerLogin(self, connInfo):

		conn2rtr   = None
		aluLogged  = False
		index      = 0

		systemIP   = connInfo['systemIP']
		deviceType = connInfo['deviceType']

		# if we have a MD-CLI device, let's make sure netmiko
		# does support it.
		deviceType = re.sub('^md_','',deviceType)

		if connInfo['useSSHTunnel'] is True:
			ip   = IP_LOCALHOST
			port = connInfo['localPort']
		else:
			ip   = connInfo['systemIP']
			port = connInfo['remotePort']

		while aluLogged == False and index < len(self.ROUTER_USER):

			tempUser = self.ROUTER_USER[index][0]
			tempPass = self.ROUTER_USER[index][1]
			index 	 = index + 1

			try:
				#conn2rtr = ConnectHandler(device_type=deviceType, host=ip, port=port, username=tempUser, password=tempPass, fast_cli=False, session_log=debug) #, log_level="DEBUG")
				if self.connInfo['sshDebug'] is True:
					debug = self.logsDirectory + "debug.debug"
					conn2rtr = ConnLogOnly(device_type=deviceType, host=ip, port=port, username=tempUser, password=tempPass, fast_cli=False, log_file=debug, log_level="DEBUG",log_format='[%(levelname)s] %(name)s: [%(threadName)s] %(message)s')
				else:
					conn2rtr = ConnLogOnly(device_type=deviceType, host=ip, port=port, username=tempUser, password=tempPass, fast_cli=False)
				aluLogged    = True
				aluLogReason = "LoggedOk"
				aluLogUser   = tempUser
				aluPass      = tempPass
			except Exception as e:
				conn2rtr     = None
				aluLogged 	 = False
				aluLogReason = str(e).replace('\n',' ').lstrip()
				aluLogUser   = tempUser
				aluPass      = "PassN/A"
				fncPrintConsole(connInfo['strConn'] + aluLogReason + ": " + systemIP)

		connInfo['conn2rtr']     = conn2rtr
		connInfo['aluLogged']    = aluLogged
		connInfo['aluLogUser']   = aluLogUser
		connInfo['aluLogReason'] = aluLogReason
		connInfo['tempPass']     = tempPass

		return connInfo

	def routerRunRoutine(self, connInfo):

		# Sending script to ALU
		tStart 		 = time.time()

		major_error_list = DICT_VENDOR[connInfo['deviceType']]['MAJOR_ERROR_LIST']
		minor_error_list = DICT_VENDOR[connInfo['deviceType']]['MINOR_ERROR_LIST']
		info_error_list  = DICT_VENDOR[connInfo['deviceType']]['INFO_ERROR_LIST']
		fin_script       = [DICT_VENDOR[connInfo['deviceType']]['FIN_SCRIPT']]

		if connInfo['cronTime']['type'] is not None:
			datos = connInfo['cronScript']
			fncPrintConsole(connInfo['strConn'] + "Establishing script with CRON...", show=1)
		else:
			datos = connInfo['pluginScript']
			fncPrintConsole(connInfo['strConn'] + "Running script per line...", show=1)

		datos = datos.split('\n')
		runStatus, aluLogReason, outRx, outRxJson = self.fncWriteToConnection(datos, connInfo)

		aluLogReason = aluLogReason.replace('\n',' ').lstrip()

		tEnd  = time.time()
		tDiff = tEnd - tStart

		## Analizing output only if writing to connection was successfull
		if aluLogReason == "":
			
			# we verify correctness of execution ...
			if any([re.compile(error, flags=re.MULTILINE).search(outRx) for error in major_error_list]):
				aluLogReason = "MajorFailed"
			elif any([re.compile(error, flags=re.MULTILINE).search(outRx) for error in minor_error_list]):
				aluLogReason = "MinorFailed"
			elif any([re.compile(error, flags=re.MULTILINE).search(outRx) for error in info_error_list]):
				aluLogReason = "InfoFailed"
			else:
				aluLogReason = "SendSuccess"

			# we verify completeness of execution ...
			if not any([re.compile(error, flags=re.MULTILINE).search(outRx) for error in fin_script]):
				aluLogReason = aluLogReason + ':Incomplete'

		fncPrintConsole(connInfo['strConn'] + "Time: " + fncFormatTime(tDiff) + ". Result: " + aluLogReason, show=1)

		connInfo['aluLogReason'] = aluLogReason
		connInfo['runStatus']    = runStatus
		connInfo['tDiff']        = tDiff
		connInfo['outRx']        = outRx
		connInfo['outRxJson']    = outRxJson

		return connInfo

	def logData(self, connInfo, logInfo, logsDirTimestamp, plugin, logsDirectory):

		writeCmd  = 'n/a'
		writeRx   = 'n/a'
		writeJson = 'n/a'

		if self.outputJob == 2:
			pluginScript = connInfo['pluginScript']
			outRx        = connInfo['outRx']
			outRxJson    = connInfo['outRxJson']			
		else:
			pluginScript = ''
			outRx        = ''
			outRxJson    = {}

		if logsDirectory:

			# Filenames
			if self.logFileName == 'hostname':
				logFname = 'hostname'
			else:
				logFname = 'systemIP'			

			aluFileCommands  = logsDirectory + connInfo[logFname] + "_commands.cfg"
			aluFileOutRx	 = logsDirectory + connInfo[logFname] + "_rx.txt"
			aluFileOutRxJson = logsDirectory + connInfo[logFname] + "_rx.json"

			if self.outputJob == 2 and connInfo['aluLogged'] == True:

				try:
					with open(aluFileOutRx,'a+') as fw:
						fw.write(outRx)
						fw.close()
						writeRx = 'yes'
				except Exception as e:
					fncPrintConsole(connInfo['strConn'] + "logData: " + str(e))
					writeRx = 'no'			

				if connInfo['cronTime']['type'] is None:

					try:
						with open(aluFileCommands,'a+') as fc:
							fc.write(pluginScript)
							fc.close()
							writeCmd = 'yes'
					except Exception as e:
						fncPrintConsole(connInfo['strConn'] + "logData: " + str(e))
						writeCmd = 'no'

				if connInfo['pluginType'] == 'show':

					if not os.path.isfile(aluFileOutRxJson):
						try:
							with open(aluFileOutRxJson,'w') as fj:
								outRxJson['name'] = connInfo['hostname']
								outRxJson['ip']   = connInfo['systemIP']
								json.dump(outRxJson,fj)
								fj.close()
								writeJson = 'yes'
						except Exception as e:
							fncPrintConsole(connInfo['strConn'] + "logData: " + str(e))					
							writeJson = 'no'
					else:
						try:
							with open(aluFileOutRxJson) as fj:
								data      = json.load(fj)
								fj.close()
							with open(aluFileOutRxJson,'w') as fj:
								outRxJson = dict(list(outRxJson.items()) + list(data.items()))
								json.dump(outRxJson,fj)
								fj.close()
							writeJson = 'yes'
						except Exception as e:
							fncPrintConsole(connInfo['strConn'] + "logData: " + str(e))
							writeJson = 'no'

		if connInfo['useSSHTunnel'] is True:
			serverName = connInfo['jumpHost']
			lenServers = len(connInfo['jumpHosts'])
		else:
			serverName = '-1'
			lenServers = '-1'

		# Building Logs ...

		logs = {
				'DateTime':logsDirTimestamp,
				'logInfo':logInfo,
				'IP':connInfo['systemIP'],
				'Timos':connInfo['timos'],
				'HostName':connInfo['hostname'],
				'HwType':connInfo['hwType'],
				'User':connInfo['username'],
				'Reason':connInfo['aluLogReason'],
				'id':str(connInfo['num']),
				'port':str(connInfo['localPort']),
				'jumpHost':serverName,
				'deviceType':connInfo['deviceType'],
				'time':float(fncFormatTime(connInfo['tDiff'], adjust=False)),
				'readTimeOut':str(connInfo['readTimeOut']),
				'servers':str(lenServers),
			}

		if self.outputJob == 2:

			logsJob2 = {
				'Plugin':plugin,
				'pluginType':connInfo['pluginType'],
				'cmdVerify':connInfo['cmdVerify'],
				'txLines':str(len(pluginScript.split('\n'))),
				'rxLines':str(len(outRx.split('\n'))),	
				'writeCmd':writeCmd,
				'writeRx':writeRx,
				'writeJson':writeJson,					
			}

			logs.update(logsJob2)

		else:

			logsJob3 = {
				'TotFtpFiles':len(connInfo['ftpFiles']),
				'TotTxFtpFiles':connInfo['ftpTotalTxFiles']
			}

			logs.update(logsJob3)			

		df = pd.DataFrame([logs])
		
		fncPrintConsole(connInfo['strConn'] + "logData: " + str(list(logs.values())))

		LOG_GLOBAL.append(df)

		connInfo['logs'] = df

		return connInfo

	def sshStop(self, connInfo):
		self.sshServer.stop()
		fncPrintConsole(connInfo['strConn'] + "SSH" + str(connInfo['num']) + " stopped ...")

	def runCron(self, connInfo):

		def setScript(cronName, script):

			cfg = ""
			cfg = cfg + f'script "{cronName}" owner "taskAutom"\nshutdown\n'
			cfg = cfg + f'location cf3:\{script}\n'
			cfg = cfg + 'no shutdown\n'
			cfg = cfg + 'exit\n'
			return cfg

		def action(cronName):

			cfg = ""
			cfg = cfg + f'action "{cronName}" owner "taskAutom"\nshutdown\n'
			cfg = cfg + 'results cf3:\\resultTestCron.txt\n'
			cfg = cfg + f'script "{cronName}" owner "taskAutom"\n'
			cfg = cfg + 'no shutdown\n'
			cfg = cfg + 'exit\n'
			return cfg

		def policy(cronName):

			cfg = ""
			cfg = cfg + f'script-policy "{cronName}" owner "taskAutom"\nshutdown\n'
			cfg = cfg + 'results cf3:\\resultTestCron.txt\n'
			cfg = cfg + f'script "{cronName}" owner "taskAutom"\n'
			cfg = cfg + 'no shutdown\n'
			cfg = cfg + 'exit\n'
			return cfg

		def schedule(connInfo):

			timos    = connInfo['timosMajor']
			cronName = connInfo['cronTime']['cronName']
			cronType = connInfo['cronTime']['type']

			cfg = ""
			cfg = cfg + f'schedule "{cronName}" owner "taskAutom"\nshutdown\n'

			if timos > 8:
				cfg = cfg + f'script-policy "{cronName}" owner "taskAutom"\n'
			else:
				cfg = cfg + f'action "{cronName}" owner "taskAutom"\n'

			if cronType == 'oneshot':

				dayOfMonth = str(connInfo['cronTime']['dayOfMonth'])
				hour       = str(connInfo['cronTime']['hour'])
				minute     = str(connInfo['cronTime']['minute'])
				month      = str(connInfo['cronTime']['month'])
				weekday    = str(connInfo['cronTime']['weekday'])
			
				cfg = cfg + 'type oneshot\n'
				cfg = cfg + f'day-of-month "{dayOfMonth}"\n'
				cfg = cfg + f'hour "{hour}"\n'
				cfg = cfg + f'minute "{minute}"\n'
				cfg = cfg + f'month "{month}"\n'
				cfg = cfg + f'weekday "{weekday}"\n'
				cfg = cfg + 'no shutdown \n'
				cfg = cfg + 'exit\n'
				cfg = cfg + 'exit all\n'
				cfg = cfg + 'admin save\n'

			elif cronType == 'periodic':

				interval = str(connInfo['cronTime']['interval'])

				cfg = cfg + 'type periodic\n'
				cfg = cfg + f'interval {interval}\n'
				cfg = cfg + 'no shutdown \n'
				cfg = cfg + 'exit all\n'
				cfg = cfg + 'admin save\n'				


			return cfg

		cronName       = connInfo['cronTime']['cronName']
		cronScriptName = connInfo['ftpRemoteFile']

		cfg = ""

		if connInfo['timosMajor'] > 8:

			cfg = cfg + "/configure system script-control\n"
			cfg = cfg + setScript(cronName, cronScriptName)
			cfg = cfg + policy(cronName)
			cfg = cfg + "/configure system cron\n"
			cfg = cfg + schedule(connInfo)

		else:

			cfg = cfg + "/configure cron\n"
			cfg = cfg + setScript(cronName, cronScriptName)
			cfg = cfg + action(cronName)
			cfg = cfg + schedule(connInfo)

		cfg = "/environment no more\n" + cfg

		connInfo['cronScript'] = cfg

		return connInfo

####################################
# Main Functions                   #
####################################

def waitBetweenRouters(dictParam):

	if dictParam['timeBetweenRouters'] > 0:
		timeToWait = dictParam['timeBetweenRouters'] / 1000
		print("Waiting " + str(timeToWait) + "s ...")
		time.sleep(timeToWait)

def createLogFolder(dictParam):

	dictParam['logsDirTimestamp'] = time.strftime('%Y-%m-%d_%H-%M-%S', time.localtime())
	if dictParam['outputJob'] in [0,2]:
		dictParam['logsDirectory'] = os.getcwd() + "/logs_" + dictParam['logsDirTimestamp'] + "_" + dictParam['logInfo'] + "_" + dictParam['pluginFileAlone'] + "/"
	else:
		dictParam['logsDirectory'] = os.getcwd() + "/logs_" + dictParam['logsDirTimestamp'] + "_" + dictParam['logInfo'] + "_ftpUpload/"
	dictParam['logsCsvFilename'] = dictParam['logsDirectory'] + "00_log.csv"

	# Verify if logsDirectory exists. If so, ask for different name ...
	if os.path.exists(dictParam['logsDirectory']):
		print("Folder " + dictParam['logsDirectory'] + " already exists.\nUse a different folder name.\nQuitting ...")
		quit()
	else:
		os.makedirs(dictParam['logsDirectory'])
		open(dictParam['logsCsvFilename'],'w').close()

	return dictParam

def getDictParam():

	parser = argparse.ArgumentParser(description='taskAutom Parameters.', prog='taskAutom', usage='%(prog)s [options]')
	parser.add_argument('-v'  ,'--version',     help='Version', action='version', version='Lucas Aimaretto - (c)2023 - laimaretto@gmail.com - Version: '+LATEST_VERSION )

	groupJobTypes = parser.add_argument_group('JobTypes')
	groupJobTypes.add_argument('-j'  ,'--jobType',       type=int, choices=[0,2,3], default=0, help='Type of job. j=0 to check data and plugin; j=2, to execute. j=3, to upload files via SCP/SFTP. When j=3, password must be entered manually. Default=0')

	groupPugin = parser.add_argument_group('Plugin')
	groupPugin.add_argument('-pt' ,'--pluginType',      type=str, help='Type of plugin.', default='show', choices=['show','config'])
	groupPugin.add_argument('-py' ,'--pluginFilename' , type=str, help='PY Template File. Optional if jobType=3.')

	groupData = parser.add_argument_group('Data Related')
	groupData.add_argument('-d'  ,'--dataFile',        type=str, required=True, help='DATA File with parameters. Either CSV or XLSX. If XLSX, enable -xls option with sheet name.')
	groupData.add_argument('-log','--logInfo' ,        type=str, required=True, help='Name of the log folder. Logs, MOP and scripts will be stored here.', )
	groupData.add_argument('-fn','--logFileName' ,     type=str, help='Name of the log fileName, either "ip" or "hostname". Default=hostname', default='hostname', choices=['ip','hostname'] )
	groupData.add_argument('-gc' ,'--dataGroupColumn', type=str, help='Only valid if using headers. Name of column, in the data file, to filter routers by. In general one should use the field where the IP of the router is. Default=ip', default='ip')
	groupData.add_argument('-uh', '--useHeader',       type=str, help='When reading data, consider first row as header. Default=yes', default='yes', choices=['no','yes'])
	groupData.add_argument('-xls' ,'--xlsSheetName',   type=str, help='Excel sheet name')
	groupData.add_argument('-so', '--strictOrder',     type=str, help='Follow strict order of routers inside the data file, row by row. If enabled, threads=1. Default=no', default='no', choices=['no','yes'])
	groupData.add_argument('-hoe','--haltOnError',     type=str, help='If using --strictOrder=yes, halts if error found on execution. Default=no', default='no', choices=['no','yes'])
	groupData.add_argument('-pbr', '--passByRow',      type=str, help='Pass data to the plugin by row (and filtered by -gc/--dataGroupColumn). Only valid with --strictOrder=no. Default=yes', default='yes', choices=['yes','no'])	

	credentialsGroup = parser.add_argument_group('Credentials')
	credentialsGroup.add_argument('-u'  ,'--username',      type=str, help='Username to connect to router.', )
	credentialsGroup.add_argument('-pf' ,'--passwordFile',  type=str, help='Filename containing the default password to access the routers. If the file contains several lines of text, only the first line will be considered as the password. Default=None', default=None)
	
	connGroup = parser.add_argument_group('Connection parameters')
	connGroup.add_argument('-th' ,'--threads' ,      type=int, help='Number of threads. Default=1', default=1,)
	connGroup.add_argument('-tun','--sshTunnel',     type=str, help='Use SSH Tunnel to routers. Default=yes', default='yes', choices=['no','yes'])
	connGroup.add_argument('-jh' ,'--jumpHostsFile', type=str, help='jumpHosts file. Default=servers.yml', default='servers.yml')
	connGroup.add_argument('-dt', '--deviceType',    type=str, help='Device Type. Default=nokia_sros', default='nokia_sros', choices=['nokia_sros','nokia_sros_telnet','nokia_srl'])
	connGroup.add_argument('-cv', '--cmdVerify',     type=str, help='Enable --cmdVerify when interacting with router. Disable only if connection problems. Default=yes', default='yes', choices=['no','yes'])
	connGroup.add_argument('-rto' ,'--readTimeOut',  type=int, help='Read Timeout. Time in seconds which to wait for data from router. Default=10', default=10,)
	connGroup.add_argument('-tbr' ,'--timeBetweenRouters',  type=int, help='Time to wait between routers, in miliseconds (ms), before sending scripts to the router. Default=0', default=0,)
	connGroup.add_argument('-axr' ,'--auxRetry',     type=int, help='Times to try obtaining aux values before "not-match". Default=10', default=10,)


	miscGroup = parser.add_argument_group('Misc')
	miscGroup.add_argument('-inv','--inventoryFile', type=str, help='Inventory file with per router connection parameters. Default=None', default=None)
	miscGroup.add_argument('-gm', '--genMop',        type=str, help='Generate MOP document in docx format. Default=no', default='no', choices=['no','yes'])
	miscGroup.add_argument('-crt','--cronTime',      type=str, nargs='+' , help='Data for CRON: type(ie: oneshot or periodic), name(ie: test).\nIf type=oneshot, need to define: month(ie april), weekday(ie monday), day-of-month(ie 28), hour(ie 17), minute(ie 45). If type=periodic, need to define: interval in seconds (ie 35).', default=[])
	miscGroup.add_argument('-sd', '--sshDebug',      type=str, help='Enables debuging of SSH interaction with the network. Stored on debug.log. Default=no', default='no', choices=['no','yes'])

	args = parser.parse_args()

	### reading parameters

	dictParam = dict(
		version            = LATEST_VERSION,		
		outputJob 		   = args.jobType,
		dataFile           = args.dataFile,
		xlsSheetName       = args.xlsSheetName,
		useHeader          = True if args.useHeader == 'yes' else False,
		passByRow          = True if args.passByRow == 'yes' else False,
		pluginFilename     = args.pluginFilename,
		username 		   = args.username,
		passwordFile       = args.passwordFile,
		password 		   = None,
		progNumThreads	   = args.threads,
		logInfo 		   = args.logInfo,
		logFileName        = args.logFileName,
		useSSHTunnel 	   = True if args.sshTunnel == 'yes' else False,
		cronTime           = args.cronTime,
		jumpHostsFile      = args.jumpHostsFile,
		genMop             = True if args.genMop == 'yes' else False,
		strictOrder        = True if args.strictOrder == 'yes' else False,
		haltOnError        = True if args.haltOnError == 'yes' else False,
		inventoryFile      = args.inventoryFile,
		deviceType         = args.deviceType,
		pluginType         = args.pluginType,
		cmdVerify          = True if args.cmdVerify == 'yes' else False,
		sshDebug           = True if args.sshDebug == 'yes' else False,
		dataGroupColumn    = args.dataGroupColumn,
		readTimeOut        = args.readTimeOut,
		timeBetweenRouters = args.timeBetweenRouters,
		auxRetry           = args.auxRetry,
	)

	################
	# Checking...

	# CronTime
	dictParam['cronTime'] = verifyCronTime(dictParam['cronTime'])
	if dictParam['cronTime']['type'] is not None:
		dictParam['pluginType']  = 'config'
		dictParam['strictOrder'] = False

	# Servers
	dictParam['jumpHosts'] = {}
	if dictParam['useSSHTunnel'] is True or dictParam['inventoryFile'] != None:
		dictParam['jumpHosts'] = verifyServers(dictParam['jumpHostsFile'])

	# Strict Order
	if dictParam['strictOrder'] is True:
		dictParam['progNumThreads'] = 1
		dictParam['passByRow'] = True

	# We verify the existence of DATA file
	dictParam['data'] = verifyData(dictParam)

	# Plugin File
	if dictParam['outputJob'] in [0,2]:

		if dictParam['pluginFilename']:
			dictParam['pluginFileAlone'] = dictParam['pluginFilename'].split('/')[-1]
			dictParam['mod'] = verifyPlugin(dictParam['pluginFilename'])
		else:
			print('Your jobType is ' + str(dictParam['outputJob']) + '. Need to specify a plugin.\nQuitting...')
			quit()	

		if not dictParam['pluginType']:
			print('Your jobType is ' + str(dictParam['outputJob']) + '. Need to specify a pluginType.\nQuitting...')
			quit()

	else:
		dictParam['mod']             = None
		dictParam['pluginFileAlone'] = None
		dictParam['pluginFilename']  = None
		dictParam['pluginType']      = None

		# If jobType = 3, the dataGroupColumn must always be 'ip'
		dictParam['dataGroupColumn'] = 'ip'

	# We check credentials
	# here we obtain the global password to be used
	# we store it in dictParam['password']
	dictParam = checkCredentials(dictParam)	

	# We obatin the list of routers to trigger connections
	# if jobType = 3, it returns a tuple (ip,fileName)
	# We generate the inventory with connections parameters for each router.
	# These default parameters can be overriden with an inventory file.
	dictParam['listOfRouters'], dictParam['inventory'] = getListOfRouters(dictParam)

	# Inventory. We update the inventory with external file.
	if dictParam['inventoryFile'] != None:
		dInv = verifyInventory(dictParam)
		for key in dInv.keys():
			dictParam['inventory'][key].update(dInv[key])

	# Check auxReytr
	if dictParam['auxRetry'] < 1:
		print('auxRetry must be greater than 0.\nQuitting...')
		quit()	

	return dictParam

def checkCredentials(dictParam):

	if dictParam['outputJob'] == 0:

		pass

	elif dictParam['outputJob'] == 2:
		
		if dictParam['username'] and dictParam['logInfo'] and (dictParam['pluginType'] or dictParam['cronTime']):
		
			if dictParam['passwordFile'] is None:

				print("\n#######################################")
				print("# About to run. Ctrl+C if not sure... #")
				print("#######################################\n")
				dictParam['password'] = getpass("### -> PASSWORD (default user: " + dictParam['username'] + "): ")

			else:
				# Trying to open the password file to obtain the password
				with open(dictParam['passwordFile']) as pf:
					dictParam['password'] = pf.readlines()[0].rstrip()
					pf.close()

		else:
			print("Your type of Job is 2, which means you are about to execute scripts on routers.\nAt least define --username, --logInfo and --pluginType.\nRun: taskAutom -h for help.\nQuitting...")
			quit()			

	elif dictParam['outputJob'] == 3:

		if dictParam['username'] and dictParam['passwordFile'] is None and dictParam['logInfo']:

			print("\n#############################################################")
			print("# You are about to do massive upload of files vía SCP/SFTP  #")
			print("# About to run. Ctrl+C if not sure...                       #")
			print("#############################################################\n")
			dictParam['password'] = getpass("### -> PASSWORD (default user: " + dictParam['username'] + "): ")

		else:
			print("Your type of Job is 3, which means you are about to send files to routers via SCP/SFTP.\nAt least define --username, --logInfo. Password must be enterd manually.\nRun: taskAutom -h for help.\nQuitting...")
			quit()	

	else:

		print("Not enough paramteres.\nAt least define --username, --logInfo, depending on the jobType.\nRun: taskAutom -h for help.\nQuitting...")
		quit()

	return dictParam

def enableLogging(dictParam):

	## Netmiko Debug
	if dictParam['sshDebug'] is True:
		
		debugFileName = dictParam['logsDirectory'] + 'debug.log'

		logger        = logging.getLogger("netmiko")
		logger.setLevel('DEBUG')
		
		handler = logging.FileHandler(debugFileName)
		handler.setLevel(logging.DEBUG)
		handler.setFormatter(logging.Formatter('[%(levelname)s] %(name)s: [%(threadName)s] %(message)s'))

		logger.addHandler(handler)

		#log_formatter = logging.Formatter("%(asctime)s [%(levelname)s] %(name)s: %(message)s [%(threadName)s] ") # I am printing thread id here
		#logging.basicConfig(filename=debugFileName, level=logging.DEBUG)
		

	return None

def fncRun(dictParam):
	"""[summary]

	Args:
		dictParam ([dict]): [Dictionary with parameters for the connections]
	Returns:
		[int]: 0
	"""

	listOfRouters = dictParam['listOfRouters']

	# We take initial time 
	timeTotalStart 	= time.time()

	# Generar threads
	threads_list = ThreadPool(dictParam['progNumThreads'])

	################
	# Running...
	if dictParam['outputJob'] == 2:

		# logInfo
		dictParam = createLogFolder(dictParam)

		# debug
		enableLogging(dictParam)

		for i, IPconnect in enumerate(listOfRouters):

			routerInfo = dictParam['inventory'][IPconnect]
			routerInfo['pluginScript'].append(renderCliLine(IPconnect, dictParam, i))

			# Wait before sending scripts to the routers ...
			waitBetweenRouters(dictParam)

			# running routine
			if dictParam['strictOrder'] is False:
				threads_list.apply_async(run_mi_thread, args=(i, routerInfo, dictParam))
			else:
				outRx = run_mi_thread(i, routerInfo, dictParam)

				if dictParam['haltOnError'] is True and outRx['aluLogReason'] not in ['SendSuccess']:
					dictParam['aluLogReason'] = outRx['aluLogReason']
					break

		if dictParam['strictOrder'] is False:
			threads_list.close()
			### The .join() implies that ALL processes/threads need to finish themselves before moving on.
			threads_list.join()

		print("all done")
		fncPrintResults(listOfRouters, timeTotalStart, dictParam)

	elif dictParam['outputJob'] == 3:

		# logInfo
		dictParam     = createLogFolder(dictParam)
		listOfRouters = list(dictParam['inventory'].keys())

		# debug
		enableLogging(dictParam)

		for i, IPconnect in enumerate(listOfRouters):

			routerInfo = dictParam['inventory'][IPconnect]		

			# Wait before sending scripts to the routers ...
			waitBetweenRouters(dictParam)		

			threads_list.apply_async(run_mi_thread, args=(i, routerInfo, dictParam))

		threads_list.close()
		threads_list.join()

		print("all done")
		fncPrintResults(listOfRouters, timeTotalStart, dictParam)

	elif dictParam['outputJob'] == 0:

		aluCliLineJob0 = ""

		# Verify if logsDirectory exists.
		if not os.path.exists(dictParam['logInfo']):
			os.makedirs(dictParam['logInfo'])		

		for i, IPconnect in enumerate(listOfRouters):

			# We firt do a rendeCli() for the router IPConnect and save the file
			tempFname = dictParam['logInfo'] + '/' + 'job0_' + IPconnect + '.cfg'
			tempCfg   = renderCliLine(IPconnect, dictParam, i=-1)

			with open(tempFname,'w') as f:
				f.write(tempCfg)
				f.close()

			# We do a second call the the renderCli() to save a global file.
			aluCliLineJob0 = aluCliLineJob0 + renderCliLine(IPconnect, dictParam, i)

		verif = verifyConfigFile(aluCliLineJob0)

		if verif != (-1,-1):
			print("\nWrong config file for router " + str(IPconnect) + "\nCheck (n,line,char): " + str(verif) + "\nQuitting...")
			quit()			

		renderMop(aluCliLineJob0, dictParam)
		fncPrintResults(listOfRouters, timeTotalStart, dictParam)

	return dictParam

def main():

	### Ready to go ...
	dictParam = getDictParam()
	fncRun(dictParam)

### To be run from the python shell
if __name__ == '__main__':
	main()